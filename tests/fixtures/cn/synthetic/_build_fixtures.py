# SPDX-License-Identifier: AGPL-3.0-only
# Copyright (c) 2025 Christopher Chen
"""Generate synthetic CN patent .docx fixtures for tw_contamination coverage.

Stage 3 deliverable: three committable docx files exercising the Q1
该等 / 该些 rejection branch in ``check_antecedent_basis_cn``. These are
the only docx fixtures that cover tw_contamination by construction —
the 10 real CN fixtures at ``tests/fixtures/cn/local/`` have zero such
findings (the contamination prefixes do not appear in real CNIPA filings
by definition).

Run from the project root::

    python tests/fixtures/cn/synthetic/_build_fixtures.py

Outputs (all committed via explicit !gitignore exceptions):

* tw_contamination_simple.docx  — 1 claim, 1 该等 reference
* tw_contamination_plural.docx  — 1 claim, 1 该些 reference
* tw_contamination_mixed.docx   — 2 claims: one with 该等, one clean

Structure mirrors ``tests/fixtures/cn/parity/create_parity_fixtures.py``
— body-anchor 五书 layout (ADR-109): section titles appear as standalone
body paragraphs, not page headers.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document


FIXTURES_DIR = Path(__file__).parent


def _add_body_anchor(doc, text: str) -> None:
    p = doc.add_paragraph(text)
    try:
        p.style = doc.styles["Heading 1"]
    except KeyError:  # pragma: no cover
        pass


def _build_docx(path: Path, title: str, claims: list[str]) -> None:
    """Assemble a minimal 五书 CN docx with the given claims.

    The specification / abstract sections carry placeholder content — the
    harness only exercises claim-side walker output, so the spec content
    is structural scaffolding to satisfy section ordering checks without
    being the focus.
    """
    doc = Document()

    # Title
    p = doc.add_paragraph(title)
    try:
        p.style = doc.styles["Title"]
    except KeyError:  # pragma: no cover
        pass

    # Abstract
    _add_body_anchor(doc, "说明书摘要")
    doc.add_paragraph(f"本发明公开了{title}，用于演示tw_contamination检测。")

    # Claims — typed prefix, no numPr
    _add_body_anchor(doc, "权利要求书")
    for claim in claims:
        doc.add_paragraph(claim)

    # Specification
    _add_body_anchor(doc, "说明书")
    doc.add_paragraph(title)
    doc.add_paragraph("技术领域")
    doc.add_paragraph(f"本发明涉及{title}。")
    doc.add_paragraph("背景技术")
    doc.add_paragraph("现有技术存在若干问题。")
    doc.add_paragraph("发明内容")
    doc.add_paragraph(f"本发明提供{title}以解决上述问题。")
    doc.add_paragraph("附图说明")
    doc.add_paragraph("图1为本发明实施例的示意图。")
    doc.add_paragraph("具体实施方式")
    doc.add_paragraph("下面结合实施例对本发明进行详细说明。")

    doc.save(str(path))


# ---------------------------------------------------------------------------
# Fixture definitions
# ---------------------------------------------------------------------------

SIMPLE_CLAIMS = [
    "1. 一种数据处理装置，包括：一处理器；以及一存储器，与所述处理器耦接，其中该等处理器配置为处理输入信号。",
]

PLURAL_CLAIMS = [
    "1. 一种通信设备，包括：一天线模块；以及一控制单元，与所述天线模块连接，其中该些天线模块用于接收信号。",
]

MIXED_CLAIMS = [
    "1. 一种成像装置，包括：一镜头；以及一图像传感器，与所述镜头光学耦合，其中该等镜头的焦距可调。",
    "2. 根据权利要求1所述的成像装置，其中所述图像传感器包括CMOS阵列。",
]


def main() -> None:
    _build_docx(
        FIXTURES_DIR / "tw_contamination_simple.docx",
        "一种数据处理装置",
        SIMPLE_CLAIMS,
    )
    _build_docx(
        FIXTURES_DIR / "tw_contamination_plural.docx",
        "一种通信设备",
        PLURAL_CLAIMS,
    )
    _build_docx(
        FIXTURES_DIR / "tw_contamination_mixed.docx",
        "一种成像装置",
        MIXED_CLAIMS,
    )
    print(f"Generated 3 synthetic tw_contamination fixtures in {FIXTURES_DIR}")


if __name__ == "__main__":
    main()

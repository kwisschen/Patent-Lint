# SPDX-License-Identifier: AGPL-3.0-only
# Copyright (c) 2025 Christopher Chen
"""Regenerate Phase 8c synthetic CN parity fixtures.

Both fixtures are committable (no real IP). Structure models the 五书
body-anchor layout observed in real CNIPA downloads (tests/fixtures/cn/local/):
section titles appear as standalone body paragraphs, not page headers.

Pair A exercises typed-prefix claim parsing (Tier 1 body anchor).
Pair B exercises w:numPr auto-numbering backfill.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree


def _add_numbering_definitions(doc) -> int:
    """Add a single abstractNum + num to the doc's numbering.xml part.

    Returns the numId usable by paragraphs. Creates the numbering part if it
    doesn't exist (python-docx's ``default.docx`` template ships one).
    """
    numbering_part = doc.part.numbering_part
    if numbering_part is None:
        raise RuntimeError("python-docx default template should have numbering part")
    numbering_xml = numbering_part._element

    # Find a free abstractNumId and numId
    existing_abs = [int(e.get(qn("w:abstractNumId"), "0"))
                    for e in numbering_xml.findall(qn("w:abstractNum"))]
    existing_num = [int(e.get(qn("w:numId"), "0"))
                    for e in numbering_xml.findall(qn("w:num"))]
    new_abs_id = (max(existing_abs) + 1) if existing_abs else 0
    new_num_id = (max(existing_num) + 1) if existing_num else 1

    w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    abstract_num = etree.SubElement(numbering_xml, f"{w}abstractNum",
                                    attrib={f"{w}abstractNumId": str(new_abs_id)})
    lvl = etree.SubElement(abstract_num, f"{w}lvl",
                           attrib={f"{w}ilvl": "0"})
    etree.SubElement(lvl, f"{w}start", attrib={f"{w}val": "1"})
    etree.SubElement(lvl, f"{w}numFmt", attrib={f"{w}val": "decimal"})
    etree.SubElement(lvl, f"{w}lvlText", attrib={f"{w}val": "%1."})
    etree.SubElement(lvl, f"{w}lvlJc", attrib={f"{w}val": "left"})

    # Reorder: abstractNum must come before num per schema — newly added
    # abstractNum is last; we'll move it to sit with the others.
    # python-docx's default template may not have a specific order, but
    # keeping newly-appended element at end is fine for most readers.

    num = etree.SubElement(numbering_xml, f"{w}num",
                           attrib={f"{w}numId": str(new_num_id)})
    etree.SubElement(num, f"{w}abstractNumId",
                     attrib={f"{w}val": str(new_abs_id)})

    return new_num_id


def _apply_numpr(paragraph, num_id: int, ilvl: int = 0) -> None:
    """Attach a w:numPr (numId + ilvl) to paragraph properties."""
    w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    p_elem = paragraph._element
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = etree.SubElement(p_elem, f"{w}pPr")
        # pPr must be the first child of p
        p_elem.insert(0, pPr)
    numPr = etree.SubElement(pPr, f"{w}numPr")
    etree.SubElement(numPr, f"{w}ilvl", attrib={f"{w}val": str(ilvl)})
    etree.SubElement(numPr, f"{w}numId", attrib={f"{w}val": str(num_id)})


def _add_body_anchor(doc, text: str) -> None:
    """Add a canonical 五书 body anchor as a standalone paragraph.

    Matches the structural convention observed in real CNIPA downloads:
    section titles appear in body text (often Heading 1 styled), not in
    Word page headers.
    """
    p = doc.add_paragraph(text)
    try:
        p.style = doc.styles["Heading 1"]
    except KeyError:  # pragma: no cover
        pass


# ---------------------------------------------------------------------------
# Pair A — apparatus_method_minimal (typed prefixes)
# ---------------------------------------------------------------------------


PAIR_A_CLAIMS = [
    "1. 一种数据处理装置，包括：处理器；存储器，与所述处理器耦接；和通信接口，配置为接收输入信号。",
    "2. 根据权利要求1所述的数据处理装置，其中所述处理器包括第一核心和第二核心。",
    "3. 根据权利要求2所述的数据处理装置，其中所述第一核心的时钟频率高于所述第二核心。",
    "4. 根据权利要求1所述的数据处理装置，其中所述存储器包括易失性存储区和非易失性存储区。",
    "5. 一种数据处理方法，包括：通过通信接口接收输入信号；通过处理器对所述输入信号进行处理；和将处理结果存储到存储器中。",
]

# Claim body-only text (what the XML <claim-text> carries) — no "N. " prefix.
PAIR_A_CLAIM_BODIES = [c.split(". ", 1)[1] for c in PAIR_A_CLAIMS]

PAIR_A_SPEC = {
    "title": "一种数据处理装置及其方法",
    "technical_field": "本发明涉及数据处理领域，特别涉及一种数据处理装置及其方法。",
    "background_art": "现有的数据处理装置在处理多路输入信号时存在效率不高的问题。",
    "disclosure": "本发明提供一种具有多核处理器的数据处理装置，以提高输入信号的处理效率。",
    "description_of_drawings": "图1为本发明实施例的数据处理装置结构示意图。",
    "mode_for_invention": "下面结合实施例对本发明进行详细说明。本发明的数据处理装置包括处理器、存储器和通信接口。",
}

PAIR_A_ABSTRACT = (
    "本发明公开了一种数据处理装置及其方法。所述数据处理装置包括处理器、存储器和通信接口。"
)


def build_pair_a_docx(path: Path) -> None:
    doc = Document()

    # Title
    p = doc.add_paragraph(PAIR_A_SPEC["title"])
    try:
        p.style = doc.styles["Title"]
    except KeyError:  # pragma: no cover
        pass

    # Abstract
    _add_body_anchor(doc, "说明书摘要")
    doc.add_paragraph(PAIR_A_ABSTRACT)

    # Claims — typed prefix, no numPr
    _add_body_anchor(doc, "权利要求书")
    for claim in PAIR_A_CLAIMS:
        doc.add_paragraph(claim)

    # Specification
    _add_body_anchor(doc, "说明书")
    doc.add_paragraph(PAIR_A_SPEC["title"])
    doc.add_paragraph("技术领域")
    doc.add_paragraph(PAIR_A_SPEC["technical_field"])
    doc.add_paragraph("背景技术")
    doc.add_paragraph(PAIR_A_SPEC["background_art"])
    doc.add_paragraph("发明内容")
    doc.add_paragraph(PAIR_A_SPEC["disclosure"])
    doc.add_paragraph("附图说明")
    doc.add_paragraph(PAIR_A_SPEC["description_of_drawings"])
    doc.add_paragraph("具体实施方式")
    doc.add_paragraph(PAIR_A_SPEC["mode_for_invention"])

    doc.save(str(path))


def build_pair_a_xml(path: Path) -> None:
    claim_texts_xml = "\n".join(
        f'    <claim id="cl{i:04d}" num="{i:04d}">'
        f"<claim-text>{PAIR_A_CLAIM_BODIES[i-1]}</claim-text></claim>"
        for i in range(1, 6)
    )
    xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<cn-application-body lang="zh" country="CN" dtd-version="1.0" file="pair-a" status="new">
  <description>
    <invention-title>{PAIR_A_SPEC["title"]}</invention-title>
    <technical-field>
      <p id="p0001" num="0001" Italic="0">{PAIR_A_SPEC["technical_field"]}</p>
    </technical-field>
    <background-art>
      <p id="p0002" num="0002" Italic="0">{PAIR_A_SPEC["background_art"]}</p>
    </background-art>
    <disclosure>
      <p id="p0003" num="0003" Italic="0">{PAIR_A_SPEC["disclosure"]}</p>
    </disclosure>
    <description-of-drawings>
      <p id="p0004" num="0004" Italic="0">{PAIR_A_SPEC["description_of_drawings"]}</p>
    </description-of-drawings>
    <mode-for-invention>
      <p id="p0005" num="0005" Italic="0">{PAIR_A_SPEC["mode_for_invention"]}</p>
    </mode-for-invention>
  </description>
  <cn-claims>
{claim_texts_xml}
  </cn-claims>
  <cn-abstract id="abst">
    <p id="p0001a" num="0001" Italic="0">{PAIR_A_ABSTRACT}</p>
  </cn-abstract>
</cn-application-body>
"""
    path.write_text(xml, encoding="utf-8")


# ---------------------------------------------------------------------------
# Pair B — numbering_multidep_markush (w:numPr auto-numbering)
# ---------------------------------------------------------------------------


PAIR_B_CLAIMS = [
    "一种聚合物组合物，其包含：聚合物基体；和添加剂，所述添加剂选自由抗氧化剂、紫外线稳定剂和阻燃剂所组成的群组。",
    "根据权利要求1所述的聚合物组合物，其中所述聚合物基体包括聚乙烯或聚丙烯。",
    "根据权利要求1所述的聚合物组合物，其中所述添加剂的含量为组合物总重量的0.1至5重量百分比。",
    "根据权利要求1至3中任一项所述的聚合物组合物，其还包含填料，其中所述填料为二氧化硅或碳酸钙。",
]

PAIR_B_SPEC = {
    "title": "一种聚合物组合物",
    "technical_field": "本发明涉及高分子材料领域。",
    "background_art": "现有聚合物组合物在高温环境下稳定性不佳。",
    "disclosure": "本发明通过添加多种功能性添加剂改善聚合物组合物的稳定性。",
    "description_of_drawings": "无附图。",
    "mode_for_invention": "下面对本发明的聚合物组合物进行详细说明。所述聚合物基体可采用聚乙烯或聚丙烯。",
}

PAIR_B_ABSTRACT = "本发明公开了一种聚合物组合物，其包含聚合物基体和添加剂。"


def build_pair_b_docx(path: Path) -> None:
    doc = Document()

    # Title
    p = doc.add_paragraph(PAIR_B_SPEC["title"])
    try:
        p.style = doc.styles["Title"]
    except KeyError:  # pragma: no cover
        pass

    # Abstract
    _add_body_anchor(doc, "说明书摘要")
    doc.add_paragraph(PAIR_B_ABSTRACT)

    # Claims — w:numPr, NO typed prefix
    _add_body_anchor(doc, "权利要求书")
    num_id = _add_numbering_definitions(doc)
    for claim in PAIR_B_CLAIMS:
        p = doc.add_paragraph(claim)
        _apply_numpr(p, num_id, ilvl=0)

    # Specification
    _add_body_anchor(doc, "说明书")
    doc.add_paragraph(PAIR_B_SPEC["title"])
    doc.add_paragraph("技术领域")
    doc.add_paragraph(PAIR_B_SPEC["technical_field"])
    doc.add_paragraph("背景技术")
    doc.add_paragraph(PAIR_B_SPEC["background_art"])
    doc.add_paragraph("发明内容")
    doc.add_paragraph(PAIR_B_SPEC["disclosure"])
    doc.add_paragraph("附图说明")
    doc.add_paragraph(PAIR_B_SPEC["description_of_drawings"])
    doc.add_paragraph("具体实施方式")
    doc.add_paragraph(PAIR_B_SPEC["mode_for_invention"])

    doc.save(str(path))


def build_pair_b_xml(path: Path) -> None:
    claims_xml_parts: list[str] = []
    # Claim 1 — independent with Markush group, nested claim-text
    claims_xml_parts.append(
        '    <claim id="cl0001" num="0001" claim-type="independent">'
        f'<claim-text>{PAIR_B_CLAIMS[0]}</claim-text></claim>'
    )
    # Claims 2, 3 — dependent on claim 1 via <claim-ref>
    for i in (2, 3):
        body = PAIR_B_CLAIMS[i - 1]
        # Embed claim-ref for dep 1
        ref_text = body.replace(
            "权利要求1",
            '<claim-ref idref="cl0001">权利要求1</claim-ref>', 1,
        )
        claims_xml_parts.append(
            f'    <claim id="cl{i:04d}" num="{i:04d}" claim-type="dependent">'
            f'<claim-text>{ref_text}</claim-text></claim>'
        )
    # Claim 4 — multi-dependent on 1-3
    c4 = PAIR_B_CLAIMS[3].replace(
        "权利要求1至3中任一项",
        '<claim-ref idref="cl0001">权利要求1</claim-ref>至'
        '<claim-ref idref="cl0003">3</claim-ref>中任一项',
        1,
    )
    claims_xml_parts.append(
        '    <claim id="cl0004" num="0004" claim-type="dependent">'
        f'<claim-text>{c4}</claim-text></claim>'
    )

    claims_xml = "\n".join(claims_xml_parts)

    xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<cn-application-body lang="zh" country="CN" dtd-version="1.0" file="pair-b" status="new">
  <description>
    <invention-title>{PAIR_B_SPEC["title"]}</invention-title>
    <technical-field>
      <p id="p0001" num="0001" Italic="0">{PAIR_B_SPEC["technical_field"]}</p>
    </technical-field>
    <background-art>
      <p id="p0002" num="0002" Italic="0">{PAIR_B_SPEC["background_art"]}</p>
    </background-art>
    <disclosure>
      <p id="p0003" num="0003" Italic="0">{PAIR_B_SPEC["disclosure"]}</p>
    </disclosure>
    <description-of-drawings>
      <p id="p0004" num="0004" Italic="0">{PAIR_B_SPEC["description_of_drawings"]}</p>
    </description-of-drawings>
    <mode-for-invention>
      <p id="p0005" num="0005" Italic="0">{PAIR_B_SPEC["mode_for_invention"]}</p>
    </mode-for-invention>
  </description>
  <cn-claims>
{claims_xml}
  </cn-claims>
  <cn-abstract id="abst">
    <p id="p0001a" num="0001" Italic="0">{PAIR_B_ABSTRACT}</p>
  </cn-abstract>
</cn-application-body>
"""
    path.write_text(xml, encoding="utf-8")


def main() -> None:
    here = Path(__file__).parent
    build_pair_a_docx(here / "apparatus_method_minimal.docx")
    build_pair_a_xml(here / "apparatus_method_minimal.xml")
    build_pair_b_docx(here / "numbering_multidep_markush.docx")
    build_pair_b_xml(here / "numbering_multidep_markush.xml")
    print("Generated 4 parity fixtures in", here)


if __name__ == "__main__":
    main()

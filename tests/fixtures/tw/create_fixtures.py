# SPDX-License-Identifier: AGPL-3.0-only
# Copyright (c) 2025 Christopher Chen
"""Generate synthetic TW patent .docx fixtures for testing."""

from __future__ import annotations

from pathlib import Path

from docx import Document


FIXTURES_DIR = Path(__file__).parent


def _add_paragraphs(doc: Document, lines: list[str]) -> None:
    for line in lines:
        doc.add_paragraph(line)


def create_invention_complete() -> None:
    """Complete invention patent with all sections, claims, symbol table."""
    doc = Document()

    lines = [
        "【發明名稱】",
        "半導體封裝結構及其製造方法",
        "【技術領域】",
        "【0001】本發明係關於一種半導體封裝結構，特別是關於一種晶片級封裝結構。",
        "【0002】隨著電子產品之微型化，半導體封裝技術亦不斷精進。",
        "【先前技術】",
        "【0003】習知之半導體封裝結構通常採用引線框架封裝技術。",
        "【0004】然而，習知技術存在散熱不佳之問題。",
        "【發明內容】",
        "【0005】本發明之目的在於提供一種改良之半導體封裝結構。",
        "【0006】本發明提供一種半導體封裝結構，包含一基板10及一晶片20。",
        "【圖式簡單說明】",
        "【0007】圖1為本發明之半導體封裝結構之剖面示意圖。",
        "【0008】圖2為本發明之半導體封裝結構之俯視示意圖。",
        "【0009】圖3為本發明之製造方法之流程圖。",
        "【實施方式】",
        "【0010】請參閱圖1，本發明之半導體封裝結構100包含一基板10。",
        "【0011】如圖2所示，基板10上設置有一晶片20。",
        "【0012】晶片20透過焊球30連接至基板10。",
        "【符號說明】",
        "10‧‧‧基板",
        "20‧‧‧晶片",
        "30‧‧‧焊球",
        "100‧‧‧半導體封裝結構",
        "【申請專利範圍】",
        "1. 一種半導體封裝結構，包含：一基板；以及一晶片，設置於該基板上。",
        "2. 如請求項1所述之半導體封裝結構，其中該基板包含一重佈線層。",
        "3. 如請求項2所述之半導體封裝結構，其中該重佈線層包含銅線路。",
        "4. 如請求項1所述之半導體封裝結構，更包含複數焊球。",
        "5. 一種半導體封裝結構之製造方法，包含：提供一基板；以及設置一晶片於該基板上。",
        "6. 如請求項5所述之製造方法，更包含形成一重佈線層。",
        "【摘要】",
        "本發明提供一種半導體封裝結構及其製造方法。該封裝結構包含基板及晶片，晶片透過焊球連接至基板上之重佈線層。",
        "【代表圖】",
        "圖1",
        "【代表圖之符號簡單說明】",
        "10‧‧‧基板",
        "20‧‧‧晶片",
    ]

    _add_paragraphs(doc, lines)
    doc.save(str(FIXTURES_DIR / "invention_complete.docx"))


def create_utility_model_complete() -> None:
    """Complete utility model patent with 新型 headers."""
    doc = Document()

    lines = [
        "【新型名稱】",
        "散熱裝置",
        "【技術領域】",
        "【0001】本新型係關於一種散熱裝置。",
        "【先前技術】",
        "【0002】習知散熱裝置之散熱效率不佳。",
        "【新型內容】",
        "【0003】本新型之目的在於提供一種改良之散熱裝置。",
        "【圖式簡單說明】",
        "【0004】圖1為本新型之散熱裝置之立體圖。",
        "【實施方式】",
        "【0005】請參閱圖1，本新型之散熱裝置包含一底座1及複數散熱鰭片2。",
        "【符號說明】",
        "1‧‧‧底座",
        "2‧‧‧散熱鰭片",
        "【申請專利範圍】",
        "1. 一種散熱裝置，包含：一底座；以及複數散熱鰭片，設置於該底座上。",
        "2. 如請求項1所述之散熱裝置，其中該底座為鋁製底座。",
        "【摘要】",
        "本新型提供一種散熱裝置。",
    ]

    _add_paragraphs(doc, lines)
    doc.save(str(FIXTURES_DIR / "utility_model_complete.docx"))


def create_missing_sections() -> None:
    """Patent missing 符號說明 and 圖式簡單說明 sections."""
    doc = Document()

    lines = [
        "【發明名稱】",
        "資料處理方法",
        "【技術領域】",
        "本發明係關於資料處理。",
        "【先前技術】",
        "習知方法效率低落。",
        "【發明內容】",
        "本發明提供一種改良之資料處理方法。",
        "【實施方式】",
        "本發明之方法包含步驟S1及步驟S2。",
        "【申請專利範圍】",
        "1. 一種資料處理方法，包含步驟S1。",
        "2. 如請求項1所述之方法，更包含步驟S2。",
        "【摘要】",
        "本發明提供一種資料處理方法。",
    ]

    _add_paragraphs(doc, lines)
    doc.save(str(FIXTURES_DIR / "missing_sections.docx"))


def create_claim_dependencies() -> None:
    """Claims with various dependency patterns including edge cases."""
    doc = Document()

    lines = [
        "【發明名稱】",
        "測試裝置",
        "【技術領域】",
        "本發明係關於測試。",
        "【發明內容】",
        "本發明提供一種測試裝置。",
        "【實施方式】",
        "本發明之測試裝置包含元件A。",
        "【申請專利範圍】",
        # Independent claim
        "1. 一種測試裝置，包含一元件A。",
        # Simple dependency: 如請求項1所述之
        "2. 如請求項1所述之測試裝置，更包含一元件B。",
        # Alternative form: 如請求項1之
        "3. 如請求項1之測試裝置，其中該元件A為金屬元件。",
        # Multi-dep with 或: 如請求項1或2之
        "4. 如請求項1或2之測試裝置，更包含一元件C。",
        # Range dep: 如請求項1~3中任一項所述之
        "5. 如請求項1~3中任一項所述之測試裝置，更包含一元件D。",
        # Second independent
        "6. 一種測試方法，包含步驟A。",
        # Self-referencing claim (error case)
        "7. 如請求項7所述之測試方法，更包含步驟B。",
        # Forward-referencing claim (error case)
        "8. 如請求項9所述之測試方法，更包含步驟C。",
        # Normal dependent on claim 6
        "9. 如請求項6所述之測試方法，更包含步驟D。",
        "【摘要】",
        "本發明提供一種測試裝置及測試方法。",
    ]

    _add_paragraphs(doc, lines)
    doc.save(str(FIXTURES_DIR / "claim_dependencies.docx"))


def create_cross_category_dependent_tw() -> None:
    """Synthetic TW fixture mirroring US cross_category_dependent.docx.

    Exercises ADR-092's immediate-parent invariant in TW vocabulary: a
    three-claim chain where each claim depends on its parent and
    introduces a new sub-component via 引用記載型式. The walker must
    resolve all references in the chain to zero findings because each
    sub-component IS introduced somewhere in the ancestor chain, just
    not at the root.
    """
    doc = Document()

    lines = [
        "【發明名稱】",
        "諧波減速裝置",
        "【技術領域】",
        "本發明係關於一種諧波減速裝置。",
        "【發明內容】",
        "本發明提供一種諧波減速裝置。",
        "【實施方式】",
        "本發明之諧波減速裝置包含波產生器、柔輪及第一剛輪。",
        "【申請專利範圍】",
        # Claim 1 — root independent, introduces 波產生器 and 柔輪
        "1. 一種諧波減速裝置，包含一波產生器及一柔輪。",
        # Claim 2 — dependent on 1, introduces 第一剛輪 (new component)
        "2. 如請求項1所述之諧波減速裝置，其中更包含一第一剛輪，所述第一剛輪與該柔輪相嚙合。",
        # Claim 3 — dependent on 2, references 第一剛輪 (defined in claim 2)
        "3. 如請求項2所述之諧波減速裝置，其中所述第一剛輪包含複數個齒狀結構。",
        # Claim 4 — dependent on 3, references 齒狀結構 (plural intro in 3)
        "4. 如請求項3所述之諧波減速裝置，其中所述齒狀結構為漸開線齒形。",
        "【摘要】",
        "本發明提供一種諧波減速裝置。",
    ]

    _add_paragraphs(doc, lines)
    doc.save(str(FIXTURES_DIR / "cross_category_dependent_tw.docx"))


def create_symbol_table_variants() -> None:
    """Symbol table with different separator styles."""
    doc = Document()

    lines = [
        "【發明名稱】",
        "連接器組件",
        "【技術領域】",
        "本發明係關於連接器。",
        "【發明內容】",
        "本發明提供一種連接器組件。",
        "【實施方式】",
        "本發明之連接器組件包含本體10。",
        "【符號說明】",
        # ‧‧‧ separator (U+2027)
        "10‧‧‧本體",
        # ... separator (ASCII dots)
        "20...端子",
        # Tab separator
        "30\t外殼",
        # Colon separator
        "40：彈片",
        # Fullwidth dot separator (U+00B7)
        "50·接觸面",
        "【申請專利範圍】",
        "1. 一種連接器組件，包含一本體。",
        "【摘要】",
        "本發明提供一種連接器組件。",
    ]

    _add_paragraphs(doc, lines)
    doc.save(str(FIXTURES_DIR / "symbol_table_variants.docx"))


if __name__ == "__main__":
    create_invention_complete()
    create_utility_model_complete()
    create_missing_sections()
    create_claim_dependencies()
    create_cross_category_dependent_tw()
    create_symbol_table_variants()
    print(f"Created fixtures in {FIXTURES_DIR}")

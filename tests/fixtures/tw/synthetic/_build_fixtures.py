# SPDX-License-Identifier: PolyForm-Noncommercial-1.0.0
# Copyright (c) 2025 Christopher Chen
"""Generate synthetic TW patent .docx fixtures for walker-mechanism coverage.

Phase A2 deliverable: 9 committable .docx fixtures covering mechanism
coverage gaps identified by Phase A1 audit (docs/tw-corpus-coverage-
audit-YYYY-MM-DD.md). These fixtures close corpus diversity gaps the
10 JP派譯版 real fixtures cannot exercise, and establish a regression
guardrail (tw_adversarial_negatives.docx) against walker over-relaxation.

Run from the project root::

    python tests/fixtures/tw/synthetic/_build_fixtures.py

Outputs (all committed via explicit !gitignore exceptions):

* tw_adversarial_negatives.docx   — MUST-FLAG claims; over-fit guardrail
* tw_copula_tiers.docx            — F12 Tier A/B/C intros (C2)
* tw_formula_reference.docx       — 式(X) expressions (C8)
* tw_markush.docx                 — 選自由...所組成之群組 (domain)
* tw_plural_prefix.docx           — TW-valid 該等/該些 (regression)
* tw_quoted_reference.docx        — 引用記載型式 beyond 具備 (8287c24 regression)
* tw_symbol_table.docx            — 符號說明 sub-figure notation (regression)
* tw_jepson_two_part.docx         — 其特徵在於/其改良在於 (domain)
* tw_dym_edge_cases.docx          — DYM quality-reject filter (B3)

Anti-over-fit authoring discipline (plan Phase A2):
1. Draft TW-first: claims conform to TIPO 專利法施行細則 §17–§21.
2. ≥2 variant drafts per mechanism across ≥2 domains.
3. Adversarial claims inline in each mechanism fixture where meaningful.
4. tw_adversarial_negatives.docx is a global guardrail: no round may
   silence its findings without explicit re-triage.
5. Build deterministically — python-docx default settings; same inputs
   produce byte-identical outputs.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document


FIXTURES_DIR = Path(__file__).parent


def _add_paragraphs(doc, lines: list[str]) -> None:
    for line in lines:
        doc.add_paragraph(line)


def _build_docx(
    path: Path,
    title: str,
    tech_field: str,
    claims: list[str],
    abstract: str,
    *,
    symbol_table: list[str] | None = None,
    drawings_desc: list[str] | None = None,
) -> None:
    """Assemble a minimal TIPO-conforming TW .docx with 【】 bracket headers.

    Spec sections beyond title / claims / abstract are scaffolding — the
    harness exercises claim-side walker output, so scaffolding satisfies
    section ordering checks without being the focus.
    """
    doc = Document()
    lines: list[str] = [
        "【發明名稱】",
        title,
        "【技術領域】",
        f"【0001】本發明係關於{tech_field}。",
        "【先前技術】",
        "【0002】習知技術存在若干問題。",
        "【發明內容】",
        f"【0003】本發明提供{title}以解決上述問題。",
        "【圖式簡單說明】",
    ]
    if drawings_desc:
        for i, desc in enumerate(drawings_desc, start=4):
            lines.append(f"【{i:04d}】{desc}")
    else:
        lines.append("【0004】圖1為本發明實施例之示意圖。")
    lines.extend([
        "【實施方式】",
        "【0010】下面結合實施例對本發明進行詳細說明。",
    ])
    if symbol_table:
        lines.append("【符號說明】")
        lines.extend(symbol_table)
    lines.append("【申請專利範圍】")
    lines.extend(claims)
    lines.extend([
        "【摘要】",
        abstract,
    ])
    _add_paragraphs(doc, lines)
    doc.save(str(path))


# ---------------------------------------------------------------------------
# Fixture 1: tw_adversarial_negatives.docx — MUST-FLAG guardrail
# ---------------------------------------------------------------------------
# Claims contain genuine missing antecedents that no F-family port should
# silence. Over-fit audit: if any C-round silences these findings, the
# mechanism is over-relaxed. Target: 3 legit_drafting_error emissions.

ADVERSARIAL_CLAIMS = [
    "1. 一種散熱裝置，包含：一基板；一晶片，設置於所述基板上；以及一連接部，連接所述基板與所述晶片。",
    "2. 如請求項1所述之散熱裝置，其中該機殼圍繞所述基板。",
    "3. 如請求項1所述之散熱裝置，其中所述導熱材填充於所述連接部。",
    "4. 如請求項1所述之散熱裝置，其中該溫控電路調節所述晶片之溫度。",
]


# ---------------------------------------------------------------------------
# Fixture 2: tw_copula_tiers.docx — F12 Tier A/B/C
# ---------------------------------------------------------------------------
# Tier A (轉變為/變為): chemistry variant
# Tier B (基於/來自): software variant
# Tier C (為/是): mechanical variant
# Each tier exercised across ≥2 claims / ≥2 domains.

COPULA_TIERS_CLAIMS = [
    "1. 一種資料處理方法，包含：接收輸入訊號；依據所述輸入訊號計算一權重；以及產生一輸出結果，其中所述權重基於所述輸入訊號產生。",
    "2. 如請求項1所述之方法，其中所述輸入訊號來自一感測器模組。",
    "3. 一種化學反應裝置，包含：一反應腔室；以及一加熱單元，設置於所述反應腔室內，其中所述反應腔室內之液體轉變為水蒸氣。",
    "4. 如請求項3所述之裝置，其中所述水蒸氣分為第一水蒸氣及第二水蒸氣。",
    "5. 一種機械結構，包含一主體及一固定件，其中所述固定件為一金屬螺栓，所述主體是一鋁合金外殼。",
    "6. 如請求項5所述之結構，其中所述金屬螺栓變為一固定狀態。",
]


# ---------------------------------------------------------------------------
# Fixture 3: tw_formula_reference.docx — 式(X) must NOT be flagged
# ---------------------------------------------------------------------------
# Claims contain scientific-notation formula references. Walker should
# suppress these in the antecedent-basis check (not referable entities).

FORMULA_REFERENCE_CLAIMS = [
    "1. 一種計算方法，包含：接收輸入資料；依據式(1)計算得到一權重值；以及依據式(2)計算得到一偏置值。",
    "2. 如請求項1所述之方法，其中所述權重值符合式(3)。",
    "3. 一種電路，包含一電阻元件及一電容元件，所述電阻元件之電阻值符合式(L-3)，所述電容元件之電容值符合式(L-4)。",
    "4. 如請求項3所述之電路，其中所述電阻值及所述電容值滿足式(R1)。",
]


# ---------------------------------------------------------------------------
# Fixture 4: tw_markush.docx — 選自由...所組成之群組
# ---------------------------------------------------------------------------
# Markush claim group. Chemistry compound + software enumeration variants.

MARKUSH_CLAIMS = [
    "1. 一種化合物，其結構選自由第一化合物、第二化合物、第三化合物及第四化合物所組成之群組。",
    "2. 如請求項1所述之化合物，其中所述第一化合物包含金屬離子。",
    "3. 一種軟體模組，其中該模組選自由語音識別模組、影像識別模組及文字識別模組所組成之群組。",
    "4. 如請求項3所述之軟體模組，其中所述語音識別模組包含一聲學模型。",
]


# ---------------------------------------------------------------------------
# Fixture 5: tw_plural_prefix.docx — TW-valid 該等/該些
# ---------------------------------------------------------------------------
# Plural reference prefixes. Walker should NOT flag as tw_contamination
# on the TW side (CN-only rejection). These are valid TIPO drafting.

PLURAL_PREFIX_CLAIMS = [
    "1. 一種感測系統，包含：複數感測器；以及複數控制器，分別耦接所述複數感測器。",
    "2. 如請求項1所述之系統，其中該等感測器為溫度感測器。",
    "3. 如請求項1所述之系統，其中該些控制器為微控制器。",
    "4. 如請求項1所述之系統，其中該等感測器耦接至該些控制器。",
]


# ---------------------------------------------------------------------------
# Fixture 6: tw_quoted_reference.docx — 引用記載型式 variants
# ---------------------------------------------------------------------------
# Regression for commit 8287c24. Claims 2/3/4 are independent (一種X
# preamble) but body-embed 如請求項1 references via 執行/根據/實施.
# Walker must treat these as independent, propagate intros via
# quoted_references field to ancestor chain.

QUOTED_REFERENCE_CLAIMS = [
    "1. 一種控制系統，包含：一處理器；以及一記憶體，耦接所述處理器，其中所述處理器存取所述記憶體之資料。",
    "2. 一種資料處理方法，執行如請求項1所述之控制系統的存取步驟，更包含：輸出所述資料至一顯示介面。",
    "3. 一種電腦可讀儲存媒體，根據如請求項1所述之控制系統運作，儲存一執行指令。",
    "4. 一種設備，實施如請求項1所述之控制系統，更包含一電源模組，供電給所述處理器。",
]


# ---------------------------------------------------------------------------
# Fixture 7: tw_symbol_table.docx — 符號說明 sub-figure notation
# ---------------------------------------------------------------------------
# Exercises check_claims_symbol_table_consistency. Claims reference
# numerals appearing in 符號說明 with sub-figure notation (圖1A, 圖2(a)).

SYMBOL_TABLE_FIXTURE = {
    "symbol_table": [
        "10‧‧‧基板（見圖1A）",
        "20‧‧‧晶片（見圖1B）",
        "30‧‧‧連接部（見圖2(a)）",
        "40‧‧‧封裝體（見圖2(b)）",
        "100‧‧‧散熱裝置",
    ],
    "drawings_desc": [
        "圖1A為本發明之基板之示意圖。",
        "圖1B為本發明之晶片之示意圖。",
        "圖2(a)為本發明之連接部之示意圖。",
        "圖2(b)為本發明之封裝體之示意圖。",
    ],
    "claims": [
        "1. 一種散熱裝置100，包含：一基板10；以及一晶片20，設置於所述基板10上。",
        "2. 如請求項1所述之散熱裝置100，更包含一連接部30，連接所述基板10與所述晶片20。",
        "3. 如請求項2所述之散熱裝置100，更包含一封裝體40，包覆所述連接部30。",
    ],
}


# ---------------------------------------------------------------------------
# Fixture 8: tw_jepson_two_part.docx — 其特徵在於/其改良在於
# ---------------------------------------------------------------------------
# Jepson-style two-part claims. Walker must recognize 其特徵在於 /
# 其改良在於 as transition phrases and scope the body accordingly.

JEPSON_CLAIMS = [
    "1. 一種資料處理系統，包含一處理器及一記憶體，所述處理器存取所述記憶體，其特徵在於，所述處理器更包含一加速單元，所述加速單元用以處理一神經網路運算。",
    "2. 如請求項1所述之系統，其中所述加速單元包含一乘加陣列。",
    "3. 一種改良的機械手臂，包含一基座及一臂部，所述臂部耦接所述基座，其改良在於，所述臂部更包含一扭力感測器，所述扭力感測器用以偵測一施加扭力。",
    "4. 如請求項3所述之機械手臂，其中所述扭力感測器設置於所述臂部之關節處。",
]


# ---------------------------------------------------------------------------
# Fixture 9: tw_dym_edge_cases.docx — DYM quality-reject filter
# ---------------------------------------------------------------------------
# B3 R21-analog test cases. Each claim pair: intro in one claim, ref
# in another with a shape the DYM filter should suppress:
#   Length-ratio:     refer to a DYM much longer than the reference
#   Leading-particle: DYM starts with 對/從/為/... or a TW reference prefix
#   Substring-wrap:   ref substring of DYM with stop-particle wrap
# All refs are genuinely missing antecedents (legit flag); the DYM
# suppression just keeps the suggestion noise down.

DYM_EDGE_CLAIMS = [
    "1. 一種控制系統，包含一控制器；以及一顯示介面，耦接所述控制器。",
    "2. 如請求項1所述之系統，其中該對應之控制器裝置輸出一狀態訊號。",
    "3. 一種量測裝置，包含一溫度感測器及一溫度讀取單元。",
    "4. 如請求項3所述之裝置，其中該測量到的溫度感測值大於一閾值。",
    "5. 一種運算設備，包含一處理器及一快取記憶體。",
    "6. 如請求項5所述之設備，其中該所述處理器執行一指令序列。",
]


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------

def main() -> None:
    _build_docx(
        FIXTURES_DIR / "tw_adversarial_negatives.docx",
        "散熱裝置",
        "一種散熱裝置，用以驗證走查器保留真實繪稿錯誤之能力",
        ADVERSARIAL_CLAIMS,
        "本發明提供一種散熱裝置，其包含基板、晶片及連接部，用以驗證走查器保留真實繪稿錯誤之能力。",
    )
    _build_docx(
        FIXTURES_DIR / "tw_copula_tiers.docx",
        "資料處理與反應裝置",
        "一種涵蓋資料處理方法、化學反應裝置及機械結構之複合實施例",
        COPULA_TIERS_CLAIMS,
        "本發明提供一種涵蓋資料處理方法、化學反應裝置及機械結構之複合實施例，以測試走查器F12繫詞模式。",
    )
    _build_docx(
        FIXTURES_DIR / "tw_formula_reference.docx",
        "計算方法與電路",
        "一種包含式(X)運算式之計算方法與電路結構",
        FORMULA_REFERENCE_CLAIMS,
        "本發明提供一種包含式(X)運算式之計算方法與電路結構，用以測試走查器對式(X)參照之處理。",
    )
    _build_docx(
        FIXTURES_DIR / "tw_markush.docx",
        "化合物與軟體模組",
        "一種涵蓋化合物與軟體模組之Markush式請求項實施例",
        MARKUSH_CLAIMS,
        "本發明提供一種涵蓋化合物與軟體模組之Markush式請求項實施例，以測試走查器對「選自由...所組成之群組」之處理。",
    )
    _build_docx(
        FIXTURES_DIR / "tw_plural_prefix.docx",
        "感測系統",
        "一種驗證該等/該些複數指代詞之感測系統",
        PLURAL_PREFIX_CLAIMS,
        "本發明提供一種感測系統，用以驗證走查器對台灣撰稿中「該等」與「該些」複數指代詞之處理。",
    )
    _build_docx(
        FIXTURES_DIR / "tw_quoted_reference.docx",
        "控制系統及相關方法",
        "一種驗證引用記載型式之控制系統及相關方法",
        QUOTED_REFERENCE_CLAIMS,
        "本發明提供一種控制系統及相關方法，用以驗證走查器對「執行/根據/實施如請求項N所述」引用記載型式之處理。",
    )
    _build_docx(
        FIXTURES_DIR / "tw_symbol_table.docx",
        "含子圖標示之散熱裝置",
        "一種使用圖1A、圖2(a)等子圖標示之散熱裝置",
        SYMBOL_TABLE_FIXTURE["claims"],
        "本發明提供一種散熱裝置，驗證走查器對符號說明中子圖標示之處理。",
        symbol_table=SYMBOL_TABLE_FIXTURE["symbol_table"],
        drawings_desc=SYMBOL_TABLE_FIXTURE["drawings_desc"],
    )
    _build_docx(
        FIXTURES_DIR / "tw_jepson_two_part.docx",
        "具加速單元之資料處理系統及改良機械手臂",
        "一種驗證其特徵在於及其改良在於兩段式請求項之實施例",
        JEPSON_CLAIMS,
        "本發明提供一種具加速單元之資料處理系統及改良機械手臂，以驗證走查器對兩段式請求項之處理。",
    )
    _build_docx(
        FIXTURES_DIR / "tw_dym_edge_cases.docx",
        "三組控制與量測裝置",
        "一種驗證DYM品質濾波之三組控制與量測裝置",
        DYM_EDGE_CLAIMS,
        "本發明提供一種包含控制系統、量測裝置及運算設備之三組實施例，用以驗證走查器Did-You-Mean品質濾波。",
    )
    print(f"Generated 9 synthetic TW fixtures in {FIXTURES_DIR}")


if __name__ == "__main__":
    main()

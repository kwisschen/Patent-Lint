# SPDX-License-Identifier: PolyForm-Noncommercial-1.0.0
# Copyright (c) 2025 Christopher Chen
"""Tests for the DOCX loader module.

Creates minimal .docx files using python-docx as test fixtures.
"""

from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from patentlint.parser.docx_loader import (
    LoadedCnDocument,
    LoadedDocument,
    LoadedTwDocument,
    load_docx,
    load_docx_cn,
    load_docx_tw,
)


def _create_simple_docx(paragraphs: list[str], path: Path) -> Path:
    """Create a simple .docx with unnumbered paragraphs."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(str(path))
    return path


def _add_numbering_to_doc(doc: Document) -> None:
    """Add a numbering definition (decimal, start=1) to the document.

    Creates the numbering XML structure that python-docx needs.
    """
    # Ensure numbering part exists
    numbering_part = doc.part.numbering_part
    numbering_xml = numbering_part._element

    # Create abstract numbering
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), "0")
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)
    abstract_num.append(lvl)
    numbering_xml.append(abstract_num)

    # Create num referencing the abstract
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), "1")
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), "0")
    num.append(abstract_num_id)
    numbering_xml.append(num)


def _set_paragraph_numbering(paragraph, num_id: str = "1", ilvl: str = "0") -> None:
    """Set numbering properties on a paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_elem = OxmlElement("w:ilvl")
    ilvl_elem.set(qn("w:val"), ilvl)
    numPr.append(ilvl_elem)
    numId_elem = OxmlElement("w:numId")
    numId_elem.set(qn("w:val"), num_id)
    numPr.append(numId_elem)
    pPr.append(numPr)


def _create_numbered_patent_docx(path: Path, spec_paragraphs: list[str],
                                  claims: list[str],
                                  abstract_text: str = "") -> Path:
    """Create a patent .docx with numbered spec paragraphs and claims."""
    doc = Document()
    _add_numbering_to_doc(doc)

    # Specification paragraphs (numbered)
    for text in spec_paragraphs:
        para = doc.add_paragraph(text)
        _set_paragraph_numbering(para)

    # Claims header
    doc.add_paragraph("CLAIMS")

    # Claim paragraphs (numbered — uses same numbering, counter resets via new numId)
    # Create a second numbering for claims
    numbering_xml = doc.part.numbering_part._element
    abstract_num2 = OxmlElement("w:abstractNum")
    abstract_num2.set(qn("w:abstractNumId"), "1")
    lvl2 = OxmlElement("w:lvl")
    lvl2.set(qn("w:ilvl"), "0")
    start2 = OxmlElement("w:start")
    start2.set(qn("w:val"), "1")
    lvl2.append(start2)
    num_fmt2 = OxmlElement("w:numFmt")
    num_fmt2.set(qn("w:val"), "decimal")
    lvl2.append(num_fmt2)
    abstract_num2.append(lvl2)
    numbering_xml.append(abstract_num2)

    num2 = OxmlElement("w:num")
    num2.set(qn("w:numId"), "2")
    abstract_ref2 = OxmlElement("w:abstractNumId")
    abstract_ref2.set(qn("w:val"), "1")
    num2.append(abstract_ref2)
    numbering_xml.append(num2)

    for text in claims:
        para = doc.add_paragraph(text)
        _set_paragraph_numbering(para, num_id="2")

    # Abstract
    if abstract_text:
        doc.add_paragraph("ABSTRACT")
        doc.add_paragraph(abstract_text)

    doc.save(str(path))
    return path


class TestLoadDocxBasics:
    """Basic loading and error handling."""

    def test_file_not_found(self):
        with pytest.raises(FileNotFoundError):
            load_docx("/nonexistent/path/file.docx")

    def test_not_docx_extension(self, tmp_path):
        txt_file = tmp_path / "test.txt"
        txt_file.write_text("hello")
        with pytest.raises(ValueError, match="Not a .docx"):
            load_docx(txt_file)

    def test_invalid_docx(self, tmp_path):
        bad_file = tmp_path / "bad.docx"
        bad_file.write_bytes(b"not a zip file")
        with pytest.raises(ValueError, match="Invalid .docx"):
            load_docx(bad_file)

    def test_simple_unnumbered_docx(self, tmp_path):
        path = tmp_path / "simple.docx"
        _create_simple_docx(["Hello world.", "Second paragraph."], path)
        result = load_docx(path)
        assert isinstance(result, LoadedDocument)
        assert "Hello world." in result.full_text
        assert result.paragraph_numberings == []
        assert result.claim_numberings == []


class TestNumberedParagraphs:
    """Tests for numbered paragraph extraction."""

    def test_spec_paragraphs_numbered(self, tmp_path):
        path = tmp_path / "patent.docx"
        _create_numbered_patent_docx(
            path,
            spec_paragraphs=[
                "This is the first paragraph of the specification.",
                "This is the second paragraph of the specification.",
            ],
            claims=["A method comprising a step."],
        )
        result = load_docx(path)
        assert result.paragraph_numberings == [1, 2]
        assert "[1]" in result.full_text
        assert "[2]" in result.full_text

    def test_claim_numberings_tracked(self, tmp_path):
        path = tmp_path / "patent.docx"
        _create_numbered_patent_docx(
            path,
            spec_paragraphs=["Specification paragraph."],
            claims=[
                "A device comprising a widget.",
                "The device of claim 1, further comprising a gadget.",
            ],
        )
        result = load_docx(path)
        assert result.claim_numberings == [1, 2]


class TestMissingEndings:
    """Tests for missing paragraph ending detection."""

    def test_missing_ending_detected(self, tmp_path):
        path = tmp_path / "patent.docx"
        _create_numbered_patent_docx(
            path,
            spec_paragraphs=[
                "This paragraph has a proper ending.",
                "This paragraph has no ending",  # Missing period
            ],
            claims=["A method comprising a step."],
        )
        result = load_docx(path)
        assert 2 in result.missing_ending_paragraphs
        assert 1 not in result.missing_ending_paragraphs

    def test_no_missing_endings(self, tmp_path):
        path = tmp_path / "patent.docx"
        _create_numbered_patent_docx(
            path,
            spec_paragraphs=[
                "First paragraph ends properly.",
                "Second paragraph also ends well!",
            ],
            claims=["A method comprising a step."],
        )
        result = load_docx(path)
        assert result.missing_ending_paragraphs == []


class TestRestrictiveWording:
    """Tests for restrictive wording detection in spec paragraphs."""

    def test_restrictive_wording_detected(self, tmp_path):
        path = tmp_path / "patent.docx"
        _create_numbered_patent_docx(
            path,
            spec_paragraphs=[
                "The feature is critical and must always operate correctly.",
                "This is a clean paragraph with no issues.",
            ],
            claims=["A method comprising a step."],
        )
        result = load_docx(path)
        assert 1 in result.improper_spec_paragraphs
        assert "critical" in result.improper_spec_phrases

    def test_no_restrictive_wording(self, tmp_path):
        path = tmp_path / "patent.docx"
        _create_numbered_patent_docx(
            path,
            spec_paragraphs=[
                "The disclosed approach provides a novel mechanism.",
                "Another clean paragraph follows here.",
            ],
            claims=["A method comprising a step."],
        )
        result = load_docx(path)
        assert result.improper_spec_paragraphs == []


class TestReturnType:
    """Tests for LoadedDocument structure."""

    def test_loaded_document_fields(self, tmp_path):
        path = tmp_path / "patent.docx"
        _create_numbered_patent_docx(
            path,
            spec_paragraphs=["A specification paragraph."],
            claims=["A method comprising a step."],
            abstract_text="A short abstract for testing purposes.",
        )
        result = load_docx(path)
        assert isinstance(result.full_text, str)
        assert isinstance(result.paragraph_numberings, list)
        assert isinstance(result.claim_numberings, list)
        assert isinstance(result.missing_ending_paragraphs, list)
        assert isinstance(result.improper_spec_paragraphs, list)
        assert isinstance(result.improper_spec_phrases, str)


class TestTrackedChanges:
    """Tests for tracked changes (revisions) detection (Bug 9)."""

    def test_clean_docx_no_tracked_changes(self, tmp_path):
        """A normal .docx without revisions -> False."""
        path = tmp_path / "clean.docx"
        _create_simple_docx(["Clean paragraph."], path)
        result = load_docx(path)
        assert result.has_tracked_changes is False

    def test_docx_with_insertion(self, tmp_path):
        """A .docx with w:ins element -> True."""
        doc = Document()
        doc.add_paragraph("Normal text.")
        # Inject a w:ins element into the body
        ins = OxmlElement("w:ins")
        ins.set(qn("w:id"), "1")
        ins.set(qn("w:author"), "Test Author")
        run = OxmlElement("w:r")
        text_elem = OxmlElement("w:t")
        text_elem.text = "inserted text"
        run.append(text_elem)
        ins.append(run)
        doc.element.body.append(ins)
        path = tmp_path / "with_ins.docx"
        doc.save(str(path))
        result = load_docx(path)
        assert result.has_tracked_changes is True

    def test_docx_with_deletion(self, tmp_path):
        """A .docx with w:del element -> True."""
        doc = Document()
        doc.add_paragraph("Normal text.")
        # Inject a w:del element into the body
        del_elem = OxmlElement("w:del")
        del_elem.set(qn("w:id"), "2")
        del_elem.set(qn("w:author"), "Test Author")
        run = OxmlElement("w:r")
        del_text = OxmlElement("w:delText")
        del_text.text = "deleted text"
        run.append(del_text)
        del_elem.append(run)
        doc.element.body.append(del_elem)
        path = tmp_path / "with_del.docx"
        doc.save(str(path))
        result = load_docx(path)
        assert result.has_tracked_changes is True

    def test_detect_tracked_changes_function(self):
        """Direct test of detect_tracked_changes()."""
        from patentlint.parser.docx_loader import detect_tracked_changes

        doc = Document()
        doc.add_paragraph("Clean.")
        assert detect_tracked_changes(doc) is False

        # Add insertion
        ins = OxmlElement("w:ins")
        ins.set(qn("w:id"), "1")
        doc.element.body.append(ins)
        assert detect_tracked_changes(doc) is True


class TestLoadDocxTwTrackedChanges:
    """Tests for tracked changes detection in load_docx_tw."""

    def test_clean_tw_docx_no_tracked_changes(self, tmp_path):
        """A normal TW .docx without revisions -> has_tracked_changes=False."""
        doc = Document()
        doc.add_paragraph("【技術領域】")
        doc.add_paragraph("本發明涉及測試。")
        path = tmp_path / "tw_clean.docx"
        doc.save(str(path))
        result = load_docx_tw(path)
        assert isinstance(result, LoadedTwDocument)
        assert result.has_tracked_changes is False
        assert len(result.paragraphs) > 0

    def test_tw_docx_with_tracked_changes(self, tmp_path):
        """A TW .docx with w:ins -> has_tracked_changes=True."""
        doc = Document()
        doc.add_paragraph("【技術領域】")
        ins = OxmlElement("w:ins")
        ins.set(qn("w:id"), "1")
        ins.set(qn("w:author"), "Test")
        doc.element.body.append(ins)
        path = tmp_path / "tw_tracked.docx"
        doc.save(str(path))
        result = load_docx_tw(path)
        assert isinstance(result, LoadedTwDocument)
        assert result.has_tracked_changes is True


class TestUnicodeNormalization:
    """Tests for invisible-whitespace normalization at extraction time.

    The loader strips ZWSP/ZWNJ/ZWJ/BOM and replaces NBSP (u00a0) with a
    regular space so downstream string matching is not silently broken by
    invisible characters pasted in from PDF or web sources.
    """

    @pytest.mark.parametrize(
        "invisible_char,name,replacement",
        [
            ("\u200b", "ZWSP", ""),
            ("\u200c", "ZWNJ", ""),
            ("\u200d", "ZWJ", ""),
            ("\ufeff", "BOM", ""),
            ("\u00a0", "NBSP", " "),
        ],
    )
    def test_load_docx_normalizes_invisible_chars(
        self, tmp_path, invisible_char, name, replacement
    ):
        spec_text = f"This{invisible_char}is the first paragraph."
        claim_text = f"A device{invisible_char}comprising a widget."
        path = tmp_path / f"patent_{name.lower()}.docx"
        _create_numbered_patent_docx(
            path,
            spec_paragraphs=[spec_text],
            claims=[claim_text],
        )
        result = load_docx(path)
        assert invisible_char not in result.full_text
        expected_spec = f"This{replacement}is the first paragraph."
        expected_claim = f"A device{replacement}comprising a widget."
        assert expected_spec in result.full_text
        assert expected_claim in result.full_text

    def test_load_docx_tw_normalizes_invisible_chars(self, tmp_path):
        # Cover all 5 chars in one TW doc.
        doc = Document()
        doc.add_paragraph("【技術領域】")
        doc.add_paragraph("本發明涉及\u200b測試與\u00a0儀器。")
        doc.add_paragraph("含有\u200c零\u200d寬\ufeff字。")
        path = tmp_path / "tw_invisible.docx"
        doc.save(str(path))
        result = load_docx_tw(path)
        joined = "\n".join(result.paragraphs)
        for ch in ("\u200b", "\u200c", "\u200d", "\ufeff"):
            assert ch not in joined
        # NBSP -> regular space
        assert "\u00a0" not in joined
        assert "涉及測試與 儀器。" in joined
        assert "含有零寬字。" in joined

    def test_load_docx_cn_normalizes_invisible_chars(self, tmp_path):
        doc = Document()
        doc.add_paragraph("技术领域")
        doc.add_paragraph("本发明涉及\u200b一种\u00a0装置。")
        path = tmp_path / "cn_invisible.docx"
        doc.save(str(path))
        result = load_docx_cn(path)
        all_text = "\n".join(p for s in result.sections for p in s.paragraphs)
        assert "\u200b" not in all_text
        assert "\u00a0" not in all_text
        assert "本发明涉及一种 装置。" in all_text


class TestLoadDocxCnTrackedChanges:
    """Tests for tracked changes detection in load_docx_cn."""

    def test_clean_cn_docx_no_tracked_changes(self, tmp_path):
        """A normal CN .docx without revisions -> has_tracked_changes=False."""
        doc = Document()
        doc.add_paragraph("技术领域")
        doc.add_paragraph("本发明涉及测试。")
        path = tmp_path / "cn_clean.docx"
        doc.save(str(path))
        result = load_docx_cn(path)
        assert isinstance(result, LoadedCnDocument)
        assert result.has_tracked_changes is False
        assert len(result.sections) > 0

    def test_cn_docx_with_tracked_changes(self, tmp_path):
        """A CN .docx with w:del -> has_tracked_changes=True."""
        doc = Document()
        doc.add_paragraph("技术领域")
        del_elem = OxmlElement("w:del")
        del_elem.set(qn("w:id"), "1")
        del_elem.set(qn("w:author"), "Test")
        doc.element.body.append(del_elem)
        path = tmp_path / "cn_tracked.docx"
        doc.save(str(path))
        result = load_docx_cn(path)
        assert isinstance(result, LoadedCnDocument)
        assert result.has_tracked_changes is True

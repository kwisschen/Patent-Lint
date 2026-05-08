# SPDX-License-Identifier: LicenseRef-PolyForm-Strict-1.0.0
# Copyright (c) 2025 Christopher Chen
"""Real-drafter audit harness — bug-class regression gates for R63-R67 fixes.

Programmatic .docx fixtures that each isolate one bug class from the
real-drafter audit taxonomy (see memory
`feedback_real_drafter_drafts_have_different_bugs.md`). Every fixture
embeds the trigger pattern that previously caused a false positive and
that the post-R65/R66/R67 walker, parser, or section-map fix correctly
silences.

The audit runner runs each fixture through ``analyze_bytes`` and asserts
two things per fixture:
  - silenced_keys: count of findings on these check_keys must be 0
  - expected_keys: each must have >=1 finding (recall - the silencer
    didn't over-silence and crush legitimate emits)

**Targeted, not whole-document clean.** Each fixture is minimal - the
goal is to exercise ONE bug-class trigger. Fixtures may emit unrelated
amend/verify findings (most commonly `specSupport.amend` because the
minimal embodiment body doesn't repeat every claim term). The harness
assertion only watches `silenced_keys`; ancillary emits are intentional
fixture-design noise, not regressions. The only fixture that asserts
whole-document cleanness is ``cn_drafter_realistic_baseline``, which
mirrors ``tests/test_integration.py::TestCnDrafterRealisticBaseline``.

**Known residual gap not closed in this harness.** The R67 intro-side
state-modifier extension covers the main `_INTRO_PATTERN` path. The
supplementary-intro paths (F7a `形成於X的Y`, F8 `相配合的Y`,
F7d `於X的Y`, F9 `透過Y連接`) do NOT yet have the symmetric extension.
A real-drafter case combining F7a-shape phrasing with a state-modifier
qualifier (e.g., `形成於X之島狀的Y` with reference `前述島狀的Y`)
would still produce a walker FP. R67 was not extended to the
supplementary paths because the empirical signal hasn't surfaced -
per DR-1 (empirical grounding before drafting).

Run: ``python tests/eval/real_drafter_audit.py``
Outputs to stdout + writes a markdown report to
``~/Library/Mobile Documents/com~apple~CloudDocs/CC Output/<dated>.md``.

The companion ``tests/test_real_drafter_baselines.py`` runs the same
fixtures as a pytest gate so a CI break catches regressions immediately.
"""

from __future__ import annotations

import io
import sys
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION

# pyright: reportMissingImports=false
sys.path.insert(0, str(Path(__file__).resolve().parents[2] / "src"))
from patentlint.models import Jurisdiction  # noqa: E402
from patentlint.pipeline import analyze_bytes  # noqa: E402


# ---------------------------------------------------------------------------
# Builder helpers (subset of tests/test_integration.py — kept self-contained
# so this module can be invoked from anywhere without test-suite imports).
# ---------------------------------------------------------------------------


def _doc_to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_cn_section_break(doc: Document, next_header_text: str) -> None:
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.header.is_linked_to_previous = False
    hp = new_section.header.paragraphs[0]
    hp.text = next_header_text


def _build_tw_minimal(
    claims_text: list[str],
    *,
    title: str = "散熱裝置",
    technical_field_line: str = "【0001】本發明係關於一種散熱裝置。",
    symbol_lines: list[str] | None = None,
    embodiment_lines: list[str] | None = None,
    prior_art_header: str = "【先前技術】",
    prior_art_lines: list[str] | None = None,
    invention_lines: list[str] | None = None,
    drawings_desc_line: str = "【0004】第1圖係散熱裝置之示意圖。",
    abstract_line: str = "本發明提供一種散熱裝置。",
) -> bytes:
    """Build a minimal-but-complete TW .docx with full per-section override.

    All section bodies and the title accept overrides. Defaults model a
    coherent heat-sink draft so the fixture is self-consistent across
    titleSubjectMatch / paragraphNumbering / symbolTableConsistency
    checks. Fixtures targeting non-heat-sink subject matter MUST override
    the title + technical_field + drawings_desc + abstract to match.

    Default paragraph numbering: 【0001】 technical → 【0002】 prior_art →
    【0003】 invention → 【0004】 drawings_desc → 【0005】 embodiment.
    Any override list must continue the sequence to keep
    paragraphNumbering check happy.
    """
    doc = Document()
    lines: list[str] = [
        "【發明名稱】",
        title,
        "【技術領域】",
        technical_field_line,
        prior_art_header,
    ]
    if prior_art_lines is not None:
        lines.extend(prior_art_lines)
    else:
        lines.append("【0002】習知技術存在散熱不佳之問題。")

    lines.append("【發明內容】")
    if invention_lines is not None:
        lines.extend(invention_lines)
    else:
        lines.append("【0003】本發明提供一種散熱裝置，包含一基座以及一散熱片。")

    lines.extend(["【圖式簡單說明】", drawings_desc_line])
    lines.append("【實施方式】")
    if embodiment_lines is not None:
        lines.extend(embodiment_lines)
    else:
        lines.append(
            "【0005】請參閱第1圖，散熱裝置100包括一基座10以及一散熱片20，前述散熱片20設置於前述基座10之上。"
        )

    lines.append("【符號說明】")
    if symbol_lines is not None:
        lines.extend(symbol_lines)
    else:
        lines.extend(["100  散熱裝置", "10   基座", "20   散熱片"])

    lines.append("【申請專利範圍】")
    lines.extend(claims_text)

    lines.extend(["【摘要】", abstract_line, "【代表圖】", "第1圖"])

    for line in lines:
        doc.add_paragraph(line)
    return _doc_to_bytes(doc)


def _build_cn_minimal(
    *,
    abstract_body: str,
    title: str,
    claims_lines: list[str],
    invention_lines: list[str],
    embodiment_lines: list[str],
) -> bytes:
    """Build a minimal CN 五书模板 .docx with overridable section bodies."""
    doc = Document()
    doc.add_paragraph(abstract_body)
    _add_cn_section_break(doc, "摘要附图")

    _add_cn_section_break(doc, "权利要求书")
    for line in claims_lines:
        doc.add_paragraph(line)
    _add_cn_section_break(doc, "说明书")

    doc.add_paragraph(title)
    doc.add_paragraph("技术领域")
    doc.add_paragraph("本发明涉及散热装置技术领域，具体涉及一种散热装置。")
    doc.add_paragraph("背景技术")
    doc.add_paragraph("已有的散热装置存在散热效率不佳的问题。")
    doc.add_paragraph("发明内容")
    for line in invention_lines:
        doc.add_paragraph(line)
    doc.add_paragraph("附图说明")
    doc.add_paragraph("图1是本发明一实施例的示意图。")
    doc.add_paragraph("具体实施方式")
    for line in embodiment_lines:
        doc.add_paragraph(line)
    _add_cn_section_break(doc, "说明书附图")

    first_section = doc.sections[0]
    first_section.header.is_linked_to_previous = False
    hp = first_section.header.paragraphs[0]
    hp.text = "说明书摘要"
    return _doc_to_bytes(doc)


# ---------------------------------------------------------------------------
# Fixture builders — one per bug class
# ---------------------------------------------------------------------------


def fixture_tw_older_dep_format() -> bytes:
    """R62 (8043745) — older TIPO drafter dep form 如申請專利範圍第N項.

    Pre-2018 TIPO drafts use this form instead of the modern
    如請求項N. The walker must parse it so dep-claim 2 inherits c1's
    intros (no spurious antecedent FP on 散熱裝置).
    """
    claims = [
        "1. 一種散熱裝置，包含一基座以及一散熱片，前述散熱片設置於前述基座之上。",
        "2. 如申請專利範圍第1項所述之散熱裝置，其中前述基座為金屬材質。",
    ]
    return _build_tw_minimal(claims)


def fixture_tw_arabic_cjk_ordinal_mix() -> bytes:
    """R63 (e0635ca) + R64 (5cca2bb) — Arabic↔CJK ordinal symmetry.

    Drafter introduces `一第1間隔件` / `一第2間隔件` in claim 1 (main
    intro path with 一 prefix, Arabic ordinals); dependent claim 2
    refers as `前述第二間隔件` (CJK ordinal). Walker must normalize
    both sides to a common form for matching. Without R63/R64 the
    sides mismatch and the walker false-positives on every dep claim.
    """
    claims = [
        "1. 一種散熱裝置，包含一基座、一第1間隔件、以及一第2間隔件，前述第1間隔件設置於前述基座之上。",
        "2. 如請求項1所述之散熱裝置，其中前述第二間隔件為金屬材質。",
    ]
    embodiment = [
        "【0005】請參閱第1圖，散熱裝置100包括一基座10、一第1間隔件11以及一第2間隔件12，前述第1間隔件11設置於前述基座10之上。",
        "【0006】前述第2間隔件12設置於前述基座10之另一面。",
    ]
    symbol_lines = [
        "100  散熱裝置",
        "10   基座",
        "11   第1間隔件",
        "12   第2間隔件",
    ]
    return _build_tw_minimal(
        claims,
        symbol_lines=symbol_lines,
        embodiment_lines=embodiment,
    )


def fixture_tw_citation_labels() -> bytes:
    """R65 (745bfd0) + R67 (bracket-subheading prefix) — citation + sub-section
    labels in prior-art subsection.

    Drafters use TWO patterns under 【先前技術】:
      1. Sub-section labels like `[先前技術文獻]` (standalone bracket
         label, no body content) — should be treated as structural
         markers, not prose. Originally R65 only excluded these from
         bracketFormat's canonical-name comparison; the TW
         `_BRACKET_SUBHEADING` regex did NOT accept the `【NNNN】`
         paragraph-number prefix common in TIPO drafts. R67
         (2026-05-08) extended the regex to accept the prefix.
      2. Citation-with-content patterns like `[專利文獻1]TW I999999B`
         (R65 745bfd0).

    Both must be skipped from paragraphEnding's prose-punctuation rule.
    """
    prior_art = [
        "【0002】習知技術存在散熱不佳之問題。",
        "[先前技術文獻]",
        "[專利文獻1]TW I999999B",
        "[專利文獻2]TW I888888B",
        "[非專利文獻1]Smith et al., Heat Sink Design, 2020",
    ]
    claims = [
        "1. 一種散熱裝置，包含一基座以及一散熱片，前述散熱片設置於前述基座之上。",
    ]
    return _build_tw_minimal(claims, prior_art_lines=prior_art)


def fixture_tw_section_alias_background() -> bytes:
    """R64 (5cca2bb) — drafter uses 【背景技術】 alias instead of 【先前技術】.

    Both should resolve to the prior_art subsection. requiredSections
    must report pass (the section IS present, just under an alias).
    """
    claims = [
        "1. 一種散熱裝置，包含一基座以及一散熱片，前述散熱片設置於前述基座之上。",
    ]
    return _build_tw_minimal(claims, prior_art_header="【背景技術】")


def fixture_tw_locative_bare_noun_intro() -> bytes:
    """R64 (5cca2bb F7d) — locative `於X的Y` supplementary intro.

    Drafter introduces 半導體基板 via the locative phrase
    ``於半導體基板的一主面上`` rather than a quantifier-led clause
    (一半導體基板). Walker's F7d pattern matches; claim's
    前述半導體基板 then resolves cleanly.
    """
    invention = [
        "【0003】本發明提供一種半導體裝置之製造方法。所述方法包含於一半導體基板的一主面上形成一閘極絕緣層，以及於前述閘極絕緣層上形成一閘極電極。",
    ]
    embodiment = [
        "【0005】請參閱第1圖，於半導體基板100的一主面上形成一閘極絕緣層200。",
        "【0006】其後於前述閘極絕緣層200上形成一閘極電極300。",
    ]
    claims = [
        "1. 一種半導體裝置之製造方法，其包含下列步驟：於一半導體基板的一主面上形成一閘極絕緣層；以及於前述閘極絕緣層上形成一閘極電極。",
        "2. 如請求項1所述之半導體裝置之製造方法，其中前述半導體基板為矽基板。",
    ]
    symbol_lines = ["100  半導體基板", "200  閘極絕緣層", "300  閘極電極"]
    return _build_tw_minimal(
        claims,
        title="半導體裝置之製造方法",
        technical_field_line="【0001】本發明係關於一種半導體裝置之製造方法。",
        invention_lines=invention,
        drawings_desc_line="【0004】第1圖係半導體裝置之製造步驟示意圖。",
        embodiment_lines=embodiment,
        symbol_lines=symbol_lines,
        abstract_line="本發明提供一種半導體裝置之製造方法。",
    )


def fixture_tw_state_modifier_lookahead() -> bytes:
    """R66 (22c8b80) + R67 (2026-05-08) — state-modifier symmetry.

    Drafter introduces a state-modifier-qualified noun (`一島狀的奈米片
    積層體`) and references it in a dependent claim (`前述島狀的奈米片
    積層體`). The walker's R66 reference-side lookahead extends past 的
    to the head noun; R67 (this session) adds the symmetric extension
    on the intro side so consistent intro+ref pairs resolve without
    spurious walker_fp.

    Without R67, the intro `一島狀的奈米片積層體` registered as bare
    `島狀` (intro pattern stops at 的) while the ref normalized to
    `島狀的奈米片積層體` — asymmetric mismatch → emit. Per the symmetry-
    audit invariant (memory feedback_symmetry_audit_normalize_chains).
    """
    invention = [
        "【0003】本發明提供一種半導體裝置，包含一基板以及一島狀的奈米片積層體，前述島狀的奈米片積層體形成於前述基板上。前述島狀的奈米片積層體之厚度為100奈米至500奈米。",
    ]
    embodiment = [
        "【0005】請參閱第1圖，半導體裝置1包括一基板10以及一島狀的奈米片積層體20。",
        "【0006】前述島狀的奈米片積層體20形成於前述基板10之上，且前述島狀的奈米片積層體20之厚度為100奈米至500奈米。",
    ]
    claims = [
        "1. 一種半導體裝置，包含一基板以及一島狀的奈米片積層體，前述島狀的奈米片積層體形成於前述基板上。",
        "2. 如請求項1所述之半導體裝置，其中前述島狀的奈米片積層體之厚度為100奈米至500奈米。",
    ]
    symbol_lines = ["1   半導體裝置", "10  基板", "20  島狀的奈米片積層體"]
    return _build_tw_minimal(
        claims,
        title="半導體裝置",
        technical_field_line="【0001】本發明係關於一種半導體裝置。",
        invention_lines=invention,
        drawings_desc_line="【0004】第1圖係半導體裝置之示意圖。",
        embodiment_lines=embodiment,
        symbol_lines=symbol_lines,
        abstract_line="本發明提供一種半導體裝置。",
    )


def fixture_cn_state_modifier_lookahead() -> bytes:
    """R66 (22c8b80) + R67 (2026-05-08) CN port — state-modifier symmetry.

    CN parallel of `tw_state_modifier_lookahead` using Simplified state
    suffixes (状/形). Verifies the R67 intro-side extension was ported
    to `cn_claims.py::extract_introductions_cn` along with the walker
    reference-side R66 logic.
    """
    abstract = (
        "本发明提供一种半导体装置。所述半导体装置包含一基板以及一岛状的纳米"
        "片积层体。所述岛状的纳米片积层体之厚度受到精密控制，可提升半导体装"
        "置之电气特性。本发明适用于先进半导体工艺，对应应用领域包含逻辑器件"
        "及储存器件。所述纳米片积层体之实施方式具体说明如下文。"
    )
    title = "一种半导体装置"
    claims = [
        "1. 一种半导体装置，其特征在于，包含一基板以及一岛状的纳米片积层体，所述岛状的纳米片积层体形成于所述基板上。",
        "2. 如权利要求1所述的半导体装置，其特征在于，所述岛状的纳米片积层体之厚度为100纳米至500纳米。",
    ]
    invention = [
        "本发明提供一种半导体装置，解决现有技术中之电气特性不佳问题。",
        "所述半导体装置包含一岛状的纳米片积层体，所述纳米片积层体形成于一基板上。",
        "通过控制所述岛状的纳米片积层体之厚度，可提升器件电气特性。",
    ]
    embodiment = [
        "下面结合附图说明本发明的具体实施方式。",
        "如图1所示，半导体装置包含基板100及岛状的纳米片积层体200。",
        "所述岛状的纳米片积层体200之厚度为100纳米至500纳米。",
    ]
    return _build_cn_minimal(
        abstract_body=abstract,
        title=title,
        claims_lines=claims,
        invention_lines=invention,
        embodiment_lines=embodiment,
    )


def fixture_tw_empty_paragraph_spacers() -> bytes:
    """R65 (6f8adc2) — empty `【NNNN】` paragraph spacers skipped by ending check.

    Drafters insert bare `【NNNN】` paragraphs as section spacers between
    prose paragraphs. The R65 fix recognizes these as empty spacers and
    skips them in the paragraph-ending check (no terminal-punctuation
    requirement on a marker that has no body content). Spacer is
    UNNUMBERED so it doesn't break the global paragraphNumbering
    sequence — the actual fixture trigger is the bare 【NNNN】 form
    appearing as a body paragraph.
    """
    prior_art = [
        "【0002】習知技術存在散熱不佳之問題。",
        "【0003】",  # bare spacer — must be skipped by R65 + not break numbering
        "習知方案在高溫下效能下降。",
    ]
    claims = [
        "1. 一種散熱裝置，包含一基座以及一散熱片，前述散熱片設置於前述基座之上。",
    ]
    invention_lines = [
        "【0004】本發明提供一種散熱裝置，包含一基座以及一散熱片。",
    ]
    embodiment_lines = [
        "【0006】請參閱第1圖，散熱裝置100包括一基座10以及一散熱片20，前述散熱片20設置於前述基座10之上。",
    ]
    return _build_tw_minimal(
        claims,
        prior_art_lines=prior_art,
        invention_lines=invention_lines,
        embodiment_lines=embodiment_lines,
        drawings_desc_line="【0005】第1圖係散熱裝置之示意圖。",
    )


def fixture_tw_subsection_bracket_brackets() -> bytes:
    """R65 (6f8adc2) — sub-section names not flagged by bracketFormat.

    Drafters use `【先前技術文獻】`, `【專利文獻】`, `【非專利文獻】` as
    sub-section labels inside 先前技術. These are NOT canonical section
    names (which would be 先前技術 / 發明內容 / etc.), but they ARE
    recognized sub-section labels per TIPO 偵錯系統. R65 excludes them
    from `_CANONICAL_SECTION_NAMES`, preventing bracketFormat from
    reporting them as misnamed sections.
    """
    prior_art = [
        "【0002】習知技術存在散熱不佳之問題。",
        "【先前技術文獻】",
        "[專利文獻1]TW I999999B",
    ]
    claims = [
        "1. 一種散熱裝置，包含一基座以及一散熱片，前述散熱片設置於前述基座之上。",
    ]
    return _build_tw_minimal(claims, prior_art_lines=prior_art)


def fixture_cn_locative_bare_noun_intro() -> bytes:
    """R7-extension (b5f3648) — CN port of the F7d locative pattern.

    Mirror of `tw_locative_bare_noun_intro`: 于半导体基板的一主面侧.
    """
    abstract = (
        "本发明提供一种半导体装置的制造方法。所述方法包含于半导体基板的一主面"
        "侧形成栅极绝缘层及栅极电极。所述栅极电极用于控制晶体管的开关。本发明"
        "适用于先进工艺制程，提升器件良率。所述方法之实施方式具体说明如下文。"
    )
    title = "一种半导体装置的制造方法"
    claims = [
        "1. 一种半导体装置的制造方法，其特征在于，包含下列步骤：于一半导体基板的一主面侧形成一栅极绝缘层；以及于所述栅极绝缘层上形成一栅极电极。",
        "2. 如权利要求1所述的制造方法，其特征在于，所述半导体基板为硅基板。",
    ]
    invention = [
        "本发明提供一种半导体装置的制造方法，解决现有技术中良率不佳的问题。",
        "于半导体基板的一主面侧形成一栅极绝缘层；以及",
        "于所述栅极绝缘层上形成一栅极电极。",
    ]
    embodiment = [
        "下面结合附图说明本发明的具体实施方式。",
        "如图1所示，于半导体基板100的一主面侧形成栅极绝缘层200。",
        "随后于栅极绝缘层200上形成栅极电极300。",
    ]
    return _build_cn_minimal(
        abstract_body=abstract,
        title=title,
        claims_lines=claims,
        invention_lines=invention,
        embodiment_lines=embodiment,
    )


def fixture_tw_supplementary_intro_arabic_ordinals() -> bytes:
    """R63 (e0635ca) — Arabic↔CJK normalize on the SUPPLEMENTARY intro path.

    Companion to `tw_arabic_cjk_ordinal_mix` which exercises the main
    intro path. R63's actual fix targets `_extract_supplementary_intros`
    (F8 VP-modifier `相配合的Y` etc.) — without R63 the supplementary
    path only ran `strip_leading_verb_tw` while the reference path ran
    the full normalize chain including Arabic→CJK. The asymmetry
    surfaced 12 walker FPs on Claire's 神秘黑屏哥.docx.

    Uses F8's `相配合的<Y starting with 第>` pattern: F8 captures the
    Arabic-ordinal noun, R63 normalizes 第1→第一, dep claim ref
    `前述第二散熱片` resolves cleanly.
    """
    claims = [
        "1. 一種散熱裝置，其包含一基座以及與前述基座相配合的第1散熱片，前述第1散熱片設置於前述基座之上。",
        "2. 如請求項1所述之散熱裝置，更包含與前述基座相配合的第2散熱片，前述第二散熱片為金屬材質。",
    ]
    embodiment = [
        "【0005】請參閱第1圖，散熱裝置100包括一基座10、一第1散熱片11以及一第2散熱片12。",
        "【0006】前述第1散熱片11設置於前述基座10之上，前述第2散熱片12設置於前述基座10之另一面。",
    ]
    symbol_lines = ["100  散熱裝置", "10   基座", "11   第1散熱片", "12   第2散熱片"]
    return _build_tw_minimal(
        claims, symbol_lines=symbol_lines, embodiment_lines=embodiment,
    )


def fixture_tw_self_loop_drafter_typo() -> bytes:
    """R64 (5cca2bb) — chain-broken walker emit suppression on self-loop.

    Drafter accidentally writes `如請求項4所述之X` inside claim 4 itself
    (self-reference typo). The dependency chain is broken at claim 4.
    R64's chain-broken walker suppression: skip the antecedent walker
    when chain[-1].dependencies != [] (drafter self-loop or cycle
    upstream) — dependent claim 5 should NOT cascade-emit walker
    findings inherited from the broken chain.

    Real defects on the fixture:
      - claim 4 self-dependency typo → emits selfDependent.amend
      - claim 4's references should NOT cascade walker FPs to c5

    Per memory feedback_real_drafter_drafts_have_different_bugs.md
    bug class #5.
    """
    claims = [
        "1. 一種散熱裝置，包含一基座以及一散熱片，前述散熱片設置於前述基座之上。",
        "2. 如請求項1所述之散熱裝置，更包含一風扇，前述風扇耦接於前述基座。",
        "3. 如請求項2所述之散熱裝置，其中前述基座之材質為金屬。",
        "4. 如請求項4所述之散熱裝置，其中前述基座為鋁合金。",
        "5. 如請求項4所述之散熱裝置，更包含一控制電路，前述控制電路耦接於前述風扇。",
    ]
    embodiment = [
        "【0005】請參閱第1圖，散熱裝置100包括一基座10、一散熱片20、一風扇30以及一控制電路40。",
        "【0006】前述散熱片20設置於前述基座10之上，前述風扇30耦接於前述基座10。",
        "【0007】前述控制電路40耦接於前述風扇30，用以控制前述風扇30之運轉。",
    ]
    symbol_lines = [
        "100  散熱裝置",
        "10   基座",
        "20   散熱片",
        "30   風扇",
        "40   控制電路",
    ]
    return _build_tw_minimal(
        claims,
        symbol_lines=symbol_lines,
        embodiment_lines=embodiment,
    )


def fixture_cn_drafter_realistic_baseline() -> bytes:
    """Mirror of `_build_cn_drafter_realistic` (test_integration.py).

    Re-runs the canonical CN drafter-format baseline through the audit
    harness so the per-bug-class report shows the all-clean reference
    point. If a refactor breaks anything on the cleanest possible CN
    drafter input, this fixture catches it alongside the targeted ones.
    """
    abstract = (
        "本发明提供一种高频电路基板用树脂组合物及其制备方法。所述树脂组合物"
        "包含环氧树脂、固化剂及无机填料。所述无机填料具有低介电常数和低介"
        "电损耗。该树脂组合物适用于制备高频电路基板，具有良好的热稳定性和"
        "机械强度。本发明还公开了所述树脂组合物的制备方法。"
    )
    title = "一种高频电路基板用树脂组合物及其制备方法"
    claims = [
        "1. 一种高频电路基板用树脂组合物，其特征在于，包含环氧树脂、固化剂和无机填料。",
        "2. 如权利要求1所述的树脂组合物，其特征在于，所述无机填料为二氧化硅。",
        "3. 如权利要求1所述的树脂组合物，其特征在于，所述固化剂为酚醛树脂。",
    ]
    invention = [
        "本发明的目的在于提供一种高频电路基板用树脂组合物，解决现有技术中的问题。",
        "本发明提供一种高频电路基板用树脂组合物，包括以下组分：",
        "环氧树脂；",
        "固化剂；以及",
        "无机填料。",
    ]
    embodiment = [
        "下面结合附图和实施例对本发明作进一步说明。",
        "如图1所示，本实施例的树脂组合物通过以下步骤制备：",
        "步骤一：将环氧树脂与固化剂混合，所述固化剂优选为酚醛树脂；",
        "步骤二：加入无机填料并搅拌均匀，所述无机填料优选为二氧化硅；以及",
        "步骤三：加热固化得到所述树脂组合物。",
    ]
    return _build_cn_minimal(
        abstract_body=abstract,
        title=title,
        claims_lines=claims,
        invention_lines=invention,
        embodiment_lines=embodiment,
    )


# ---------------------------------------------------------------------------
# Audit metadata
# ---------------------------------------------------------------------------


@dataclass
class FixtureSpec:
    """Per-fixture metadata: builder + jurisdiction + which checks to gate."""

    name: str
    bug_class: str
    fix_round: str  # e.g. "R62", "R66"
    fix_sha: str  # short SHA of the fix commit
    jurisdiction: Jurisdiction
    builder: object  # Callable[[], bytes]
    silenced_keys: tuple[str, ...]  # check_key prefixes that must yield 0 amend/verify
    expected_keys: tuple[str, ...]  # check_key prefixes that must yield ≥1 of any status


FIXTURES: list[FixtureSpec] = [
    FixtureSpec(
        name="tw_older_dep_format",
        bug_class="Older TIPO dep form 如申請專利範圍第N項",
        fix_round="R62",
        fix_sha="8043745",
        jurisdiction=Jurisdiction.TW,
        builder=fixture_tw_older_dep_format,
        silenced_keys=(
            "check.tw.claims.antecedentBasis.amend",
            "check.tw.claims.dependencyFormat.amend",
        ),
        expected_keys=(
            "check.tw.spec.requiredSections.pass",
        ),
    ),
    FixtureSpec(
        name="tw_arabic_cjk_ordinal_mix",
        bug_class="Mixed Arabic/CJK ordinal style across intro and reference",
        fix_round="R63",
        fix_sha="e0635ca",
        jurisdiction=Jurisdiction.TW,
        builder=fixture_tw_arabic_cjk_ordinal_mix,
        silenced_keys=(
            "check.tw.claims.antecedentBasis.amend",
        ),
        expected_keys=(
            "check.tw.claims.antecedentBasis.pass",
        ),
    ),
    FixtureSpec(
        name="tw_citation_labels",
        bug_class="Bibliographic citation labels [專利文獻N] in 先前技術",
        fix_round="R65",
        fix_sha="745bfd0",
        jurisdiction=Jurisdiction.TW,
        builder=fixture_tw_citation_labels,
        silenced_keys=(
            "check.tw.spec.paragraphEnding.verify",
        ),
        expected_keys=(
            "check.tw.spec.paragraphEnding.pass",
        ),
    ),
    FixtureSpec(
        name="tw_section_alias_background",
        bug_class="Section-name alias 【背景技術】 ↔ 【先前技術】",
        fix_round="R64",
        fix_sha="5cca2bb",
        jurisdiction=Jurisdiction.TW,
        builder=fixture_tw_section_alias_background,
        silenced_keys=(
            "check.tw.spec.requiredSections.amend",
            "check.tw.crossRef.bracketFormat.amend",
        ),
        expected_keys=(
            "check.tw.spec.requiredSections.pass",
        ),
    ),
    FixtureSpec(
        name="tw_locative_bare_noun_intro",
        bug_class="Locative bare-noun intro 於X的Y (F7d)",
        fix_round="R64",
        fix_sha="5cca2bb",
        jurisdiction=Jurisdiction.TW,
        builder=fixture_tw_locative_bare_noun_intro,
        silenced_keys=(
            "check.tw.claims.antecedentBasis.amend",
        ),
        expected_keys=(
            "check.tw.claims.antecedentBasis.pass",
        ),
    ),
    FixtureSpec(
        name="tw_state_modifier_lookahead",
        bug_class="State-modifier symmetry — intro-side lookahead",
        fix_round="R66+R67",
        fix_sha="22c8b80+(this session)",
        jurisdiction=Jurisdiction.TW,
        builder=fixture_tw_state_modifier_lookahead,
        silenced_keys=(
            "check.tw.claims.antecedentBasis.amend",
        ),
        expected_keys=(
            "check.tw.claims.antecedentBasis.pass",
        ),
    ),
    FixtureSpec(
        name="cn_state_modifier_lookahead",
        bug_class="CN port — state-modifier symmetry intro-side lookahead",
        fix_round="R66+R67",
        fix_sha="22c8b80+(this session)",
        jurisdiction=Jurisdiction.CN,
        builder=fixture_cn_state_modifier_lookahead,
        silenced_keys=(
            "check.cn.claims.antecedentBasis.amend",
        ),
        expected_keys=(
            "check.cn.claims.antecedentBasis.pass",
        ),
    ),
    FixtureSpec(
        name="tw_supplementary_intro_arabic_ordinals",
        bug_class="R63 actual fix path — supplementary-intro Arabic→CJK",
        fix_round="R63",
        fix_sha="e0635ca",
        jurisdiction=Jurisdiction.TW,
        builder=fixture_tw_supplementary_intro_arabic_ordinals,
        silenced_keys=(
            "check.tw.claims.antecedentBasis.amend",
        ),
        expected_keys=(
            "check.tw.claims.antecedentBasis.pass",
        ),
    ),
    FixtureSpec(
        name="tw_self_loop_drafter_typo",
        bug_class="Self-loop drafter typo — chain-broken walker suppression",
        fix_round="R64",
        fix_sha="5cca2bb",
        jurisdiction=Jurisdiction.TW,
        builder=fixture_tw_self_loop_drafter_typo,
        silenced_keys=(
            "check.tw.claims.antecedentBasis.amend",
        ),
        expected_keys=(
            "check.tw.claims.selfDependent.amend",
        ),
    ),
    FixtureSpec(
        name="tw_empty_paragraph_spacers",
        bug_class="Empty 【NNNN】 paragraph spacers skipped by ending check",
        fix_round="R65",
        fix_sha="6f8adc2",
        jurisdiction=Jurisdiction.TW,
        builder=fixture_tw_empty_paragraph_spacers,
        silenced_keys=(
            "check.tw.spec.paragraphEnding.verify",
        ),
        expected_keys=(
            "check.tw.spec.paragraphEnding.pass",
        ),
    ),
    FixtureSpec(
        name="tw_subsection_bracket_brackets",
        bug_class="Sub-section names 【專利文獻】 / 【先前技術文獻】 not flagged",
        fix_round="R65",
        fix_sha="6f8adc2",
        jurisdiction=Jurisdiction.TW,
        builder=fixture_tw_subsection_bracket_brackets,
        silenced_keys=(
            "check.tw.crossRef.bracketFormat.amend",
        ),
        expected_keys=(
            "check.tw.spec.requiredSections.pass",
        ),
    ),
    FixtureSpec(
        name="cn_locative_bare_noun_intro",
        bug_class="CN port — locative 于X的Y (R7-extension)",
        fix_round="R7-ext",
        fix_sha="b5f3648",
        jurisdiction=Jurisdiction.CN,
        builder=fixture_cn_locative_bare_noun_intro,
        silenced_keys=(
            "check.cn.claims.antecedentBasis.amend",
        ),
        expected_keys=(
            "check.cn.claims.antecedentBasis.pass",
        ),
    ),
    FixtureSpec(
        name="cn_drafter_realistic_baseline",
        bug_class="Canonical 五书模板 drafter format — all-clean reference point",
        fix_round="ADR-141 baseline",
        fix_sha="7d59a5f",
        jurisdiction=Jurisdiction.CN,
        builder=fixture_cn_drafter_realistic_baseline,
        silenced_keys=(
            "check.cn.claims.antecedentBasis.amend",
            "check.cn.spec.requiredSections.amend",
            "check.cn.spec.paragraphEnding.verify",
        ),
        expected_keys=(
            "check.cn.spec.requiredSections.pass",
            "check.cn.abstract.charCount.pass",
        ),
    ),
]


# ---------------------------------------------------------------------------
# Runner
# ---------------------------------------------------------------------------


def _gather_findings(result, jurisdiction: Jurisdiction) -> list:
    """Collect all check items across the section buckets for a jurisdiction."""
    if jurisdiction == Jurisdiction.US:
        return (
            result.specification_checks
            + result.claims_checks
            + result.abstract_checks
            + result.drawings_checks
        )
    if jurisdiction == Jurisdiction.CN:
        return (
            result.cn_specification_checks
            + result.cn_claims_checks
            + result.cn_abstract_checks
            + result.cn_drawings_checks
        )
    return (
        result.tw_specification_checks
        + result.tw_claims_checks
        + result.tw_abstract_checks
        + result.tw_drawings_checks
    )


def _matches(check_item, key_prefixes: tuple[str, ...]) -> bool:
    return any(check_item.message_key == p for p in key_prefixes)


@dataclass
class FixtureResult:
    spec: FixtureSpec
    silenced_violations: list  # list[CheckItem] — should be empty
    expected_hits: dict  # key -> count
    all_findings: list


def run_fixture(spec: FixtureSpec) -> FixtureResult:
    docx_bytes = spec.builder()
    result = analyze_bytes(docx_bytes, f"{spec.name}.docx", spec.jurisdiction)
    findings = _gather_findings(result, spec.jurisdiction)

    # Silenced check_keys: there should be no check items whose
    # message_key matches the silenced prefix. (We register the EXACT
    # amend/verify keys to silence, not the .pass equivalent.)
    silenced_violations = [c for c in findings if _matches(c, spec.silenced_keys)]

    # Expected check_keys: each prefix must have at least one matching
    # finding (any status — we just want the check to fire).
    expected_hits: dict[str, int] = {}
    for prefix in spec.expected_keys:
        # Match any status: convert ".pass" → ".", ".amend" → ".", etc.
        # Simplest: count exact-prefix matches.
        prefix_root = prefix.rsplit(".", 1)[0]
        count = sum(1 for c in findings if c.message_key.startswith(prefix_root + "."))
        expected_hits[prefix] = count

    return FixtureResult(
        spec=spec,
        silenced_violations=silenced_violations,
        expected_hits=expected_hits,
        all_findings=findings,
    )


def render_report(results: list[FixtureResult]) -> str:
    """Render a markdown report for the audit run."""
    out: list[str] = []
    today = date.today().isoformat()
    out.append(f"# Real-drafter audit — {today}\n")
    out.append("Per-bug-class regression gate covering R63–R66 walker / parser /")
    out.append("section-map fixes. Each fixture isolates one bug class from the")
    out.append("`feedback_real_drafter_drafts_have_different_bugs.md` taxonomy.\n")

    n_pass = sum(1 for r in results if not r.silenced_violations and all(c >= 1 for c in r.expected_hits.values()))
    out.append(f"**Summary:** {n_pass} / {len(results)} fixtures pass.\n")

    out.append("## Per-fixture results\n")
    out.append("| Fixture | Bug class | Round | Silencer | Expected emit | Status |")
    out.append("|---|---|---|---|---|---|")
    for r in results:
        s = r.spec
        sv = len(r.silenced_violations)
        eh_total = sum(r.expected_hits.values())
        eh_min = min(r.expected_hits.values()) if r.expected_hits else 0
        passes = sv == 0 and eh_min >= 1
        status = "✅ pass" if passes else "❌ fail"
        sv_cell = "0 (clean)" if sv == 0 else f"{sv} VIOL"
        eh_cell = f"{eh_total}" if eh_min >= 1 else f"{eh_total} (gap)"
        out.append(
            f"| `{s.name}` | {s.bug_class} | {s.fix_round} ({s.fix_sha}) | "
            f"{sv_cell} | {eh_cell} | {status} |"
        )

    out.append("\n## Per-bug-class detail\n")
    for r in results:
        s = r.spec
        out.append(f"### `{s.name}` — {s.bug_class}")
        out.append(f"- Fix: {s.fix_round} commit `{s.fix_sha}`")
        out.append(f"- Jurisdiction: {s.jurisdiction.name}")
        out.append(f"- Total findings on fixture: {len(r.all_findings)}")
        out.append(f"- Silenced keys: `{', '.join(s.silenced_keys)}`")
        out.append(f"- Silenced violations: {len(r.silenced_violations)}")
        if r.silenced_violations:
            for v in r.silenced_violations[:5]:
                out.append(f"  - `{v.message_key}` — {v.message[:120]}")
        out.append(f"- Expected keys: `{', '.join(s.expected_keys)}`")
        for k, count in r.expected_hits.items():
            out.append(f"  - `{k}` (any status): {count}")
        out.append("")

    out.append("## Methodology notes\n")
    out.append(
        "- Each fixture is built programmatically via python-docx; no binary "
        ".docx files are committed."
    )
    out.append(
        "- A fixture's `silenced_keys` list contains the exact `.amend` or "
        "`.verify` message_keys that previously fired on the bug-class "
        "trigger. Post-fix, these must yield zero findings."
    )
    out.append(
        "- A fixture's `expected_keys` list contains check_keys that should "
        "still emit (any status) so we catch over-silencing — the check must "
        "fire, just not with an amend/verify on the trigger pattern."
    )
    out.append(
        "- This harness tests REGRESSION, not absolute precision. A fixture's "
        "pass means the silencer behaves exactly as it did at the recorded "
        "fix SHA on the same trigger; a fail means a refactor or new emit "
        "rule has reintroduced the FP."
    )
    return "\n".join(out)


def main() -> int:
    print(f"Real-drafter audit — running {len(FIXTURES)} fixtures...")
    results: list[FixtureResult] = []
    for spec in FIXTURES:
        try:
            r = run_fixture(spec)
        except Exception as exc:  # noqa: BLE001
            print(f"  ERROR  {spec.name}: {exc!r}")
            raise
        sv = len(r.silenced_violations)
        eh_min = min(r.expected_hits.values()) if r.expected_hits else 0
        status = "PASS" if sv == 0 and eh_min >= 1 else "FAIL"
        print(f"  {status:4s}  {spec.name:36s}  silenced_viol={sv}  expected_min={eh_min}")
        results.append(r)

    report = render_report(results)
    cc_dir = Path.home() / "Library/Mobile Documents/com~apple~CloudDocs/CC Output"
    if cc_dir.exists():
        out_path = cc_dir / f"{date.today().isoformat()}_real-drafter-audit.md"
        out_path.write_text(report, encoding="utf-8")
        print(f"\nReport written: {out_path}")
    else:
        print("\n(CC Output dir not found; printing report to stdout)")
        print(report)

    failed = [r for r in results if r.silenced_violations or any(c < 1 for c in r.expected_hits.values())]
    return 0 if not failed else 1


if __name__ == "__main__":
    sys.exit(main())

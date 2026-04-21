# SPDX-License-Identifier: PolyForm-Noncommercial-1.0.0
# Copyright (c) 2025 Christopher Chen
"""Tests for the FastAPI application.

Uses FastAPI TestClient for endpoint testing with real .docx fixtures.
"""

import io

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fastapi.testclient import TestClient

from patentlint.api.app import app


client = TestClient(app)


def _add_numbering_to_doc(doc: Document) -> None:
    """Add a decimal numbering definition starting at 1."""
    numbering_part = doc.part.numbering_part
    numbering_xml = numbering_part._element

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), "0")
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)
    abstract_num.append(lvl)
    numbering_xml.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), "1")
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), "0")
    num.append(abstract_num_id)
    numbering_xml.append(num)


def _set_paragraph_numbering(paragraph, num_id: str = "1") -> None:
    """Set numbering properties on a paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_elem = OxmlElement("w:ilvl")
    ilvl_elem.set(qn("w:val"), "0")
    numPr.append(ilvl_elem)
    numId_elem = OxmlElement("w:numId")
    numId_elem.set(qn("w:val"), num_id)
    numPr.append(numId_elem)
    pPr.append(numPr)


def _create_patent_docx_bytes(spec_texts: list[str] | None = None,
                               claims_texts: list[str] | None = None,
                               abstract_text: str = "") -> bytes:
    """Create a patent .docx in memory and return its bytes."""
    doc = Document()

    if spec_texts or claims_texts:
        _add_numbering_to_doc(doc)

        # Add second numbering for claims
        numbering_xml = doc.part.numbering_part._element
        abstract_num2 = OxmlElement("w:abstractNum")
        abstract_num2.set(qn("w:abstractNumId"), "1")
        lvl2 = OxmlElement("w:lvl")
        lvl2.set(qn("w:ilvl"), "0")
        start2 = OxmlElement("w:start")
        start2.set(qn("w:val"), "1")
        lvl2.append(start2)
        num_fmt2 = OxmlElement("w:numFmt")
        num_fmt2.set(qn("w:val"), "decimal")
        lvl2.append(num_fmt2)
        abstract_num2.append(lvl2)
        numbering_xml.append(abstract_num2)
        num2 = OxmlElement("w:num")
        num2.set(qn("w:numId"), "2")
        abstract_ref2 = OxmlElement("w:abstractNumId")
        abstract_ref2.set(qn("w:val"), "1")
        num2.append(abstract_ref2)
        numbering_xml.append(num2)

    if spec_texts:
        for text in spec_texts:
            para = doc.add_paragraph(text)
            _set_paragraph_numbering(para, num_id="1")

    if claims_texts:
        doc.add_paragraph("CLAIMS")
        for text in claims_texts:
            para = doc.add_paragraph(text)
            _set_paragraph_numbering(para, num_id="2")

    if abstract_text:
        doc.add_paragraph("ABSTRACT")
        doc.add_paragraph(abstract_text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestHealthEndpoint:

    def test_health_returns_ok(self):
        response = client.get("/api/health")
        assert response.status_code == 200
        data = response.json()
        assert data["status"] == "ok"
        assert data["version"] == "1.0.0"


class TestAnalyzeEndpoint:

    def test_reject_non_docx(self):
        response = client.post(
            "/api/analyze",
            files={"file": ("test.txt", b"hello", "text/plain")},
        )
        assert response.status_code == 400
        assert "docx" in response.json()["detail"].lower()

    def test_reject_invalid_docx(self):
        response = client.post(
            "/api/analyze",
            files={"file": ("test.docx", b"not a real docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert response.status_code == 400

    def test_analyze_minimal_docx(self):
        docx_bytes = _create_patent_docx_bytes(
            spec_texts=["This is a specification paragraph."],
            claims_texts=["A method comprising a step of processing data."],
            abstract_text="A method for processing data is described.",
        )
        response = client.post(
            "/api/analyze",
            files={"file": ("patent.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert response.status_code == 200
        data = response.json()
        assert "claims" in data
        assert "abstract_word_count" in data
        assert "paragraphs_sequential" in data
        assert "figures_count" in data

    def test_analyze_returns_claim_data(self):
        docx_bytes = _create_patent_docx_bytes(
            spec_texts=["Specification paragraph."],
            claims_texts=[
                "A device comprising a widget.",
                "The device of claim 1, further comprising a gadget.",
            ],
        )
        response = client.post(
            "/api/analyze",
            files={"file": ("patent.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert response.status_code == 200
        data = response.json()
        assert data["independent_claims_count"] >= 0
        assert data["dependent_claims_count"] >= 0

    def test_analyze_empty_docx(self):
        """An empty .docx should return results with zero counts, not error."""
        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        response = client.post(
            "/api/analyze",
            files={"file": ("empty.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert response.status_code == 200
        data = response.json()
        assert data["claims"] == []
        assert data["abstract_word_count"] == 0


class TestAnalyzeReportEndpoint:

    def test_returns_pdf(self):
        docx_bytes = _create_patent_docx_bytes(
            spec_texts=["This is a specification paragraph."],
            claims_texts=["A method comprising a step of processing data."],
            abstract_text="A method for processing data is described.",
        )
        response = client.post(
            "/api/analyze/report",
            files={"file": ("test.docx", docx_bytes,
                   "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert response.status_code == 200
        assert response.headers["content-type"] == "application/pdf"
        assert response.content[:4] == b"%PDF"

    def test_rejects_non_docx(self):
        response = client.post(
            "/api/analyze/report",
            files={"file": ("test.txt", b"not a docx", "text/plain")},
        )
        assert response.status_code == 400


class TestCORS:

    def test_cors_headers_present(self):
        response = client.options(
            "/api/health",
            headers={"Origin": "http://localhost:3000", "Access-Control-Request-Method": "GET"},
        )
        assert response.status_code == 200
        assert "access-control-allow-origin" in response.headers


class TestPickLocale:
    """Locale resolution precedence: ?locale= > Accept-Language > en."""

    def test_query_param_wins(self):
        from patentlint.api.app import _pick_locale

        assert _pick_locale("zh-TW", "en-US") == "zh-TW"

    def test_accept_language_fallback(self):
        from patentlint.api.app import _pick_locale

        assert _pick_locale(None, "zh-TW,zh;q=0.9,en;q=0.8") == "zh-TW"

    def test_default_en(self):
        from patentlint.api.app import _pick_locale

        assert _pick_locale(None, None) == "en"

    def test_bcp47_normalization(self):
        """zh-Hant-TW → zh-TW, zh-Hans → zh-CN, en-US → en."""
        from patentlint.api.app import _pick_locale

        assert _pick_locale(None, "zh-Hant-TW") == "zh-TW"
        assert _pick_locale(None, "zh-Hans,zh;q=0.8") == "zh-CN"
        assert _pick_locale(None, "en-US") == "en"
        assert _pick_locale(None, "ja-JP") == "ja"

    def test_unsupported_locale_falls_back(self):
        from patentlint.api.app import _pick_locale

        assert _pick_locale(None, "fr-FR") == "en"
        assert _pick_locale("xx-YY", None) == "en"


class TestReportEndpointLocale:

    def test_accept_language_produces_localized_pdf(self):
        docx_bytes = _create_patent_docx_bytes(
            spec_texts=["Specification."],
            claims_texts=["A device."],
            abstract_text="Abstract.",
        )
        response = client.post(
            "/api/analyze/report",
            files={"file": ("test.docx", docx_bytes,
                   "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            headers={"Accept-Language": "zh-TW"},
        )
        assert response.status_code == 200
        assert response.content[:4] == b"%PDF"
        # PDFs compress text — length sanity check only. Substring
        # assertions on localized text would require a PDF parser.
        assert len(response.content) > 1000

    def test_query_param_overrides_accept_language(self):
        """?locale=en should win even if Accept-Language says zh-TW."""
        from unittest.mock import patch

        docx_bytes = _create_patent_docx_bytes(spec_texts=["Spec."])
        with patch(
            "patentlint.api.app.render_pdf", return_value=b"%PDF-1.4 stub"
        ) as mock_render:
            client.post(
                "/api/analyze/report?locale=en",
                files={"file": ("test.docx", docx_bytes,
                       "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                headers={"Accept-Language": "zh-TW,zh;q=0.9"},
            )
        _, kwargs = mock_render.call_args
        assert kwargs.get("locale") == "en"

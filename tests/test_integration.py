# SPDX-License-Identifier: PolyForm-Noncommercial-1.0.0
# Copyright (c) 2025 Christopher Chen
"""End-to-end integration tests calling analyze_bytes() with full-structure .docx fixtures.

Each fixture is built programmatically with python-docx. No binary .docx files are checked in.
"""

from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from patentlint.models import Jurisdiction
from patentlint.pipeline import analyze_bytes


# ---------------------------------------------------------------------------
# Helper: add Word section break + set next section header
# ---------------------------------------------------------------------------

def _add_cn_section_break(doc: Document, next_header_text: str) -> None:
    """Insert a Word section break (new page) and set the header of the new section.

    CN 五書模板 uses Word section headers (黑体 16pt centered) to delimit
    the five patent documents. python-docx's add_section() inserts a
    <w:sectPr> in the last body paragraph's <w:pPr>, which is exactly what
    load_docx_cn() looks for when grouping paragraphs.
    """
    from docx.enum.section import WD_SECTION
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.header.is_linked_to_previous = False
    hp = new_section.header.paragraphs[0]
    hp.text = next_header_text


def _add_us_numbering(doc: Document) -> None:
    """Add Word numbering definitions so paragraphs get sequential numbers.

    Creates two numbering lists:
    - numId 1: spec paragraphs (decimal, start=1)
    - numId 2: claims (decimal, start=1)
    """
    numbering_part = doc.part.numbering_part
    numbering_xml = numbering_part._element

    for abstract_id, num_id in [("0", "1"), ("1", "2")]:
        abstract_num = OxmlElement("w:abstractNum")
        abstract_num.set(qn("w:abstractNumId"), abstract_id)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal")
        lvl.append(num_fmt)
        abstract_num.append(lvl)
        numbering_xml.append(abstract_num)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), num_id)
        ref = OxmlElement("w:abstractNumId")
        ref.set(qn("w:val"), abstract_id)
        num.append(ref)
        numbering_xml.append(num)


def _set_para_num(paragraph, num_id: str = "1") -> None:
    """Set Word numbering on a paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numPr.append(ilvl)
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), num_id)
    numPr.append(nid)
    pPr.append(numPr)


def _add_us_numbering_gap_id(doc: Document, abstract_id: str, num_id: str, start_at: int) -> None:
    """Append an additional numbering definition with a non-default startAt.

    Used to build fixtures with a paragraph-numbering gap: existing numId=1
    runs 1..N contiguously; this one starts at ``start_at`` so paragraphs
    switching to ``num_id`` produce a discontinuity in ``paragraph_numberings``.
    """
    numbering_part = doc.part.numbering_part
    numbering_xml = numbering_part._element

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), abstract_id)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), str(start_at))
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)
    abstract_num.append(lvl)
    numbering_xml.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), abstract_id)
    num.append(ref)
    numbering_xml.append(num)


def _inject_tracked_insertion(paragraph, text: str = "inserted revision") -> None:
    """Inject a w:ins tracked-change element into a paragraph's XML.

    python-docx has no public Track Changes API, so we build the OOXML
    element directly. ``detect_tracked_changes`` in docx_loader.py fires
    when any w:ins or w:del element is present anywhere in the body.
    """
    ins_elem = OxmlElement("w:ins")
    ins_elem.set(qn("w:id"), "1")
    ins_elem.set(qn("w:author"), "Test Author")
    ins_elem.set(qn("w:date"), "2026-04-19T00:00:00Z")
    run = OxmlElement("w:r")
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)
    ins_elem.append(run)
    paragraph._element.append(ins_elem)


def _doc_to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Fixture builders — TW
# ---------------------------------------------------------------------------

def _build_tw_invention_all_pass() -> bytes:
    """Clean TW invention patent — zero AMEND/VERIFY expected."""
    doc = Document()
    lines = [
        "【發明名稱】",
        "散熱裝置",
        "【技術領域】",
        "【0001】本發明係關於一種散熱裝置，特別是關於一種用於電子元件之散熱裝置。",
        "【先前技術】",
        "【0002】隨著電子產業的發展，電子元件的散熱問題日益重要。",
        "【0003】習知散熱裝置包括散熱鰭片及風扇，然而其散熱效率有限。",
        "【發明內容】",
        "【0004】本發明之目的在於提供一種散熱裝置，以提高散熱效率。",
        "【0005】本發明提供一種散熱裝置，包括一基座及複數散熱鰭片。",
        "【圖式簡單說明】",
        "【0006】第1圖係本發明散熱裝置之立體示意圖。",
        "【0007】第2圖係本發明散熱裝置之剖面示意圖。",
        "【0008】第3圖係本發明散熱裝置之分解示意圖。",
        "【實施方式】",
        "【0009】請參閱第1圖，本發明散熱裝置100包括一基座10及複數散熱鰭片20。",
        "【0010】基座10具有一頂面及一底面，底面用於接觸電子元件。",
        "【0011】請參閱第2圖，散熱鰭片20設置於基座10之頂面。",
        "【0012】請參閱第3圖，散熱裝置100更包括一風扇30。",
        "【符號說明】",
        "100  散熱裝置",
        "10   基座",
        "20   散熱鰭片",
        "30   風扇",
        "【申請專利範圍】",
        "1. 一種散熱裝置，包括：",
        "一基座(10)，具有一頂面及一底面；及",
        "複數散熱鰭片(20)，設置於該基座(10)之該頂面；",
        "其特徵在於，該基座(10)之該底面具有一導熱層。",
        "2. 如請求項1所述之散熱裝置，其中該導熱層之材料為銅。",
        "3. 如請求項1所述之散熱裝置，更包括一風扇(30)，設置於該複數散熱鰭片(20)之上方。",
        "4. 如請求項3所述之散熱裝置，其中該風扇(30)之轉速為可調節的。",
        "5. 一種散熱系統，包括：",
        "一電子元件；及",
        "一如請求項1所述之散熱裝置，設置於該電子元件上；",
        "其特徵在於，該散熱系統更包括一溫度感測器。",
        "【摘要】",
        "本發明提供一種散熱裝置，包括一基座及複數散熱鰭片，基座之底面具有導熱層，以提高散熱效率。",
        "【代表圖】",
        "第1圖",
    ]
    for line in lines:
        doc.add_paragraph(line)
    return _doc_to_bytes(doc)


def _build_tw_utility_model_all_pass() -> bytes:
    """Clean TW utility model patent — zero AMEND/VERIFY expected."""
    doc = Document()
    lines = [
        "【新型名稱】",
        "散熱鰭片結構改良",
        "【技術領域】",
        "【0001】本新型係關於一種散熱鰭片結構改良，特別是關於一種改良之散熱鰭片。",
        "【先前技術】",
        "【0002】習知散熱鰭片結構存在散熱效率不佳之問題。",
        "【新型內容】",
        "【0003】本新型之目的在於提供一種散熱鰭片結構改良，以提高散熱效率。",
        "【圖式簡單說明】",
        "【0004】第1圖係本新型散熱鰭片結構改良之立體示意圖。",
        "【0005】第2圖係本新型散熱鰭片結構改良之剖面示意圖。",
        "【實施方式】",
        "【0006】請參閱第1圖，本新型散熱鰭片結構改良200包括一基板40及複數鰭片50。",
        "【0007】請參閱第2圖，鰭片50設置於基板40之表面。",
        "【符號說明】",
        "200  散熱鰭片結構改良",
        "40   基板",
        "50   鰭片",
        "【申請專利範圍】",
        "1. 一種散熱鰭片結構改良，包括：",
        "一基板(40)；及",
        "複數鰭片(50)，設置於該基板(40)之表面；",
        "其改良在於，該複數鰭片(50)呈波浪狀排列。",
        "2. 如請求項1所述之散熱鰭片結構改良，其中該基板(40)之材料為鋁合金。",
        "3. 如請求項1所述之散熱鰭片結構改良，其中該複數鰭片(50)之高度為5至10毫米。",
        "【摘要】",
        "本新型提供一種散熱鰭片結構改良，包括一基板及複數鰭片，鰭片呈波浪狀排列以提高散熱效率。",
        "【代表圖】",
        "第1圖",
    ]
    for line in lines:
        doc.add_paragraph(line)
    return _doc_to_bytes(doc)


def _build_tw_invention_multi_fail() -> bytes:
    """TW invention patent with intentional errors to trigger multiple checks."""
    doc = Document()
    lines = [
        "【發明名稱】",
        "測試裝置",
        "【技術領域】",
        "【0001】本新型係關於一種測試裝置。",  # 本新型 in invention → patentTypeTerminology
        "【先前技術】",
        "【0002】習知技術存在問題。",
        # Gap: skips 0003 → paragraphNumbering AMEND
        "【發明內容】",
        "【0004】本發明提供一種測試裝置。",
        "【圖式簡單說明】",
        "【0005】第1圖係測試裝置之示意圖。",
        "【實施方式】",
        "【0006】請參閱第1圖，測試裝置300包括一組件310及一模組320。",
        "【0007】組件310連接於模組320",  # No ending punctuation → paragraphEnding
        # Missing 【符號說明】→ symbolTablePresence AMEND
        "【申請專利範圍】",
        "1. 一種測試裝置，包括：",
        "一組件(310)；及",
        "一模組(320)，連接於該組件(310)；",
        "其特徵在於，該模組(320)具有一控制單元。",
        "2. 如請求項1所述之測試裝置，其中該控制單元為微處理器。",
        "3. 如請求項4所述之測試裝置，其中該微處理器之頻率為1GHz。",  # forward dep
        "4. 如請求項1所述之測試裝置，更包括一感測器。",
        "【摘要】",
        # Over 250 chars + contains 最佳
        "本發明提供一種測試裝置，包括一組件及一模組，模組具有控制單元，"
        "本發明之測試裝置為最佳的散熱解決方案，能夠有效降低電子元件之溫度，"
        "提高系統之穩定性與可靠性，適用於各種電子產品之散熱需求，"
        "包括手機、平板電腦、筆記型電腦、桌上型電腦、伺服器等電子設備。",
        "【代表圖】",
        "第1圖",
    ]
    for line in lines:
        doc.add_paragraph(line)
    return _doc_to_bytes(doc)


def _build_cn_docx_all_pass() -> bytes:
    """Clean CN patent in 五書模板 .docx format — zero AMEND/VERIFY expected."""
    doc = Document()

    # Section 1: 说明书摘要
    doc.add_paragraph("本发明提供一种数据处理方法，包括数据预处理步骤和特征提取步骤，"
                      "特征提取步骤采用深度学习模型，以提高数据分类的准确率。")
    _add_cn_section_break(doc, "摘要附图")

    # Section 2: 摘要附图 (empty)
    _add_cn_section_break(doc, "权利要求书")

    # Section 3: 权利要求书
    doc.add_paragraph("1. 一种数据处理方法，包括：")
    doc.add_paragraph("数据预处理步骤，用于对原始数据进行清洗；及")
    doc.add_paragraph("特征提取步骤，用于从清洗后的数据中提取特征向量；")
    doc.add_paragraph("其特征在于，所述特征提取步骤采用深度学习模型。")
    doc.add_paragraph("2. 如权利要求1所述的数据处理方法，其中所述深度学习模型为卷积神经网络。")
    doc.add_paragraph("3. 如权利要求2所述的数据处理方法，其中所述卷积神经网络包括至少三个卷积层。")
    doc.add_paragraph("4. 如权利要求1所述的数据处理方法，还包括模型训练步骤。")
    doc.add_paragraph("5. 如权利要求4所述的数据处理方法，其中所述模型训练步骤采用反向传播算法。")
    doc.add_paragraph("6. 一种数据处理装置，包括处理器和存储器，所述处理器执行如权利要求1所述的数据处理方法。")
    _add_cn_section_break(doc, "说明书")

    # Section 4: 说明书 (body with sub-section headers)
    doc.add_paragraph("一种数据处理方法")  # title
    doc.add_paragraph("技术领域")
    doc.add_paragraph("本发明涉及一种数据处理方法，特别涉及一种基于机器学习的数据分类方法。")
    doc.add_paragraph("背景技术")
    doc.add_paragraph("随着大数据时代的到来，数据分类技术在各领域得到广泛应用。")
    doc.add_paragraph("现有数据分类方法主要包括决策树、支持向量机等，但在处理高维数据时效率较低。")
    doc.add_paragraph("发明内容")
    doc.add_paragraph("本发明的目的在于提供一种数据处理方法，以提高数据分类的准确率和效率。")
    doc.add_paragraph("本发明提供一种数据处理方法，包括数据预处理步骤和特征提取步骤。")
    doc.add_paragraph("附图说明")
    doc.add_paragraph("图1是本发明数据处理方法的流程示意图。")
    doc.add_paragraph("图2是本发明数据分类模型的结构示意图。")
    doc.add_paragraph("具体实施方式")
    doc.add_paragraph("请参阅图1，本发明数据处理方法包括以下步骤。")
    doc.add_paragraph("步骤一，对原始数据进行预处理，去除噪声数据。")
    doc.add_paragraph("请参阅图2，将预处理后的数据输入至分类模型中。")
    _add_cn_section_break(doc, "说明书附图")

    # Section 5: 说明书附图 (empty)

    # Set the first section's header to 说明书摘要
    first_section = doc.sections[0]
    first_section.header.is_linked_to_previous = False
    hp = first_section.header.paragraphs[0]
    hp.text = "说明书摘要"

    return _doc_to_bytes(doc)


def _build_cn_docx_multi_fail() -> bytes:
    """CN patent with intentional errors to trigger multiple checks."""
    doc = Document()

    # Section 1: 说明书摘要 — over 300 chars
    doc.add_paragraph(
        "本发明提供一种数据处理方法，包括数据预处理步骤和特征提取步骤，"
        "特征提取步骤采用深度学习模型，以提高数据分类的准确率和效率，"
        "该方法能够有效处理大规模高维数据集，实现快速准确的数据分类，"
        "适用于图像识别、自然语言处理、语音识别等多种人工智能应用场景，"
        "具有广泛的应用前景和重要的实用价值，能够显著提升系统性能。"
    )
    _add_cn_section_break(doc, "摘要附图")

    # Section 2: 摘要附图 (empty)
    _add_cn_section_break(doc, "权利要求书")

    # Section 3: 权利要求书 — gap (skips 4) + 请求项 terminology
    doc.add_paragraph("1. 一种数据处理方法，包括：")
    doc.add_paragraph("数据预处理步骤；及特征提取步骤；")
    doc.add_paragraph("其特征在于，所述特征提取步骤采用深度学习模型。")
    doc.add_paragraph("2. 如权利要求1所述的数据处理方法，其中所述深度学习模型为卷积神经网络。")
    doc.add_paragraph("3. 如权利要求1所述的数据处理方法，还包括模型训练步骤。")
    # Skips claim 4
    doc.add_paragraph("5. 如请求项3所述的数据处理方法，其中所述模型训练步骤采用反向传播算法。")
    _add_cn_section_break(doc, "说明书")

    # Section 4: 说明书 — missing 背景技术 + paragraph numbering + bad ending
    doc.add_paragraph("一种数据处理方法")  # title
    doc.add_paragraph("技术领域")
    doc.add_paragraph("[0001]本发明涉及数据处理方法。")  # user-added numbering
    # Missing 背景技术 section
    doc.add_paragraph("发明内容")
    doc.add_paragraph("[0002]本发明提供一种数据处理方法。")
    doc.add_paragraph("附图说明")
    doc.add_paragraph("图1是本发明的流程图。")
    doc.add_paragraph("具体实施方式")
    doc.add_paragraph("请参阅图1，本发明包括以下步骤")  # no period ending
    _add_cn_section_break(doc, "说明书附图")

    # Section 5: 说明书附图 (empty)

    # Set first section header
    first_section = doc.sections[0]
    first_section.header.is_linked_to_previous = False
    hp = first_section.header.paragraphs[0]
    hp.text = "说明书摘要"

    return _doc_to_bytes(doc)


def _build_us_full_length() -> bytes:
    """Realistic-length US patent specification with 10 claims."""
    doc = Document()
    _add_us_numbering(doc)

    # Section headers are unnumbered; body paragraphs are numbered.
    # This matches real patent .docx structure where only body text is numbered.
    # Headers and title are unnumbered; the title text itself is also unnumbered
    headers = {
        "TITLE OF THE INVENTION",
        "Thermal Management System for Electronic Devices",
        "CROSS-REFERENCE TO RELATED APPLICATIONS",
        "FIELD OF THE INVENTION",
        "BACKGROUND OF THE INVENTION",
        "SUMMARY OF THE INVENTION",
        "BRIEF DESCRIPTION OF THE DRAWINGS",
        "DETAILED DESCRIPTION OF THE PREFERRED EMBODIMENTS",
    }
    spec_paragraphs = [
        "TITLE OF THE INVENTION",
        "Thermal Management System for Electronic Devices",
        "CROSS-REFERENCE TO RELATED APPLICATIONS",
        "This application claims benefit of U.S. Provisional Application "
        "No. 63/123,456, filed December 10, 2023.",
        "FIELD OF THE INVENTION",
        "The present invention relates to a thermal management system for "
        "electronic devices.",
        "BACKGROUND OF THE INVENTION",
        "Electronic devices generate heat during operation. Conventional heat sinks "
        "include finned structures attached to device surfaces. However, these "
        "approaches have limited thermal dissipation capacity as device densities increase.",
        "U.S. Pat. No. 9,876,543 discloses a heat sink with copper fins. U.S. Patent "
        "Application Publication No. 2022/0123456 describes a liquid cooling system.",
        "SUMMARY OF THE INVENTION",
        "In accordance with an aspect of the present invention, a thermal management "
        "system includes a base plate, a plurality of heat dissipation fins, and a "
        "phase-change material layer disposed between the base plate and the fins.",
        "BRIEF DESCRIPTION OF THE DRAWINGS",
        "FIG. 1 is a perspective view of the thermal management system.",
        "FIG. 2 is a cross-sectional view taken along line A-A of FIG. 1.",
        "FIG. 3 is an exploded view of the thermal management system.",
        "FIG. 4 is a flowchart of a method of manufacturing the thermal management system.",
        "FIG. 5 is a graph showing thermal performance comparison.",
        "DETAILED DESCRIPTION OF THE PREFERRED EMBODIMENTS",
        "Referring to FIG. 1, a thermal management system 100 includes a base plate 10, "
        "a plurality of heat dissipation fins 20, and a phase-change material layer 30. "
        "The base plate 10 has a top surface 12 and a bottom surface 14.",
        "Referring to FIG. 2, the phase-change material layer 30 is disposed between the "
        "base plate 10 and the heat dissipation fins 20. The phase-change material layer "
        "30 comprises a paraffin-based compound.",
        "Referring to FIG. 3, the heat dissipation fins 20 are detachably coupled to the "
        "base plate 10 via a plurality of fastening elements 40.",
        "Referring to FIG. 4, a method of manufacturing the thermal management system 100 "
        "includes forming the base plate 10, applying the phase-change material layer 30, "
        "and attaching the heat dissipation fins 20.",
        "Referring to FIG. 5, the thermal management system 100 achieves a 40% improvement "
        "in thermal dissipation compared to conventional approaches.",
    ]
    for text in spec_paragraphs:
        para = doc.add_paragraph(text)
        if text not in headers:
            _set_para_num(para, "1")

    # Claims
    doc.add_paragraph("CLAIMS")

    claims = [
        "A thermal management system comprising: "
        "a base plate having a top surface and a bottom surface; "
        "a plurality of heat dissipation fins disposed on the top surface of the base plate; and "
        "a phase-change material layer disposed between the base plate and the plurality of heat dissipation fins.",
        "The thermal management system of claim 1, wherein the base plate comprises copper.",
        "The thermal management system of claim 1, wherein the phase-change material layer "
        "comprises a paraffin-based compound.",
        "The thermal management system of claim 1, further comprising a plurality of fastening "
        "elements coupling the plurality of heat dissipation fins to the base plate.",
        "The thermal management system of claim 4, wherein the plurality of fastening elements "
        "are snap-fit connectors.",
        "The thermal management system of claim 1, wherein the plurality of heat dissipation "
        "fins are arranged in a radial pattern.",
        "A method of manufacturing a thermal management system, the method comprising: "
        "forming a base plate having a top surface and a bottom surface; "
        "applying a phase-change material layer on the top surface of the base plate; and "
        "attaching a plurality of heat dissipation fins to the phase-change material layer.",
        "The method of claim 7, further comprising curing the phase-change material layer "
        "at a temperature between 60 and 80 degrees Celsius.",
        "The method of claim 7, wherein attaching the plurality of heat dissipation fins "
        "comprises snap-fitting the fins to the base plate.",
        "A thermal management system comprising: "
        "a base plate; "
        "a heat dissipation assembly disposed on the base plate, the heat dissipation assembly "
        "comprising a plurality of fins and a phase-change material; and "
        "a temperature sensor disposed on the base plate, "
        "wherein the temperature sensor is configured to monitor a temperature of the base plate.",
    ]
    for text in claims:
        para = doc.add_paragraph(text)
        _set_para_num(para, "2")

    doc.add_paragraph("ABSTRACT")
    doc.add_paragraph(
        "A thermal management system includes a base plate, a plurality of heat dissipation "
        "fins, and a phase-change material layer disposed between the base plate and the fins. "
        "The phase-change material layer provides enhanced thermal dissipation by absorbing and "
        "releasing latent heat during phase transitions. A method of manufacturing the system "
        "is also disclosed."
    )

    return _doc_to_bytes(doc)


def _build_us_cluster1_defects() -> bytes:
    """Engineered US patent exercising 6 zero-coverage defect checks.

    Phase E cluster 1: one testspec hits CRM non-transitory, Jepson,
    Markush open-transition, omnibus, wherein-comma, and extra-period
    checks. Each claim is isolated so any single claim triggers exactly
    one of the six target message_keys without cross-contaminating the
    other five.
    """
    doc = Document()
    _add_us_numbering(doc)

    headers = {
        "TITLE OF THE INVENTION",
        "Defective Claims Testbed for Phase E Coverage",
        "FIELD OF THE INVENTION",
        "BACKGROUND OF THE INVENTION",
        "SUMMARY OF THE INVENTION",
        "BRIEF DESCRIPTION OF THE DRAWINGS",
        "DETAILED DESCRIPTION OF THE PREFERRED EMBODIMENTS",
    }
    spec_paragraphs = [
        "TITLE OF THE INVENTION",
        "Defective Claims Testbed for Phase E Coverage",
        "FIELD OF THE INVENTION",
        "The present invention relates to engineered defect coverage for US claim checks.",
        "BACKGROUND OF THE INVENTION",
        "Conventional test suites lack synthetic fixtures that exercise each defect class in isolation.",
        "SUMMARY OF THE INVENTION",
        "A testbed document recites six claims, each engineered to trigger a specific defect class.",
        "BRIEF DESCRIPTION OF THE DRAWINGS",
        "FIG. 1 is a schematic overview of the testbed.",
        "DETAILED DESCRIPTION OF THE PREFERRED EMBODIMENTS",
        "Referring to FIG. 1, the testbed document 100 includes a specification section and a claims section.",
    ]
    for text in spec_paragraphs:
        para = doc.add_paragraph(text)
        if text not in headers:
            _set_para_num(para, "1")

    doc.add_paragraph("CLAIMS")

    claims = [
        # Claim 1: CRM without non-transitory qualifier
        "A computer-readable storage medium comprising instructions that, "
        "when executed by a processor, cause the processor to receive sensor data, "
        "process the sensor data to identify an anomaly, and transmit an alert to a monitoring system.",
        # Claim 2: Jepson format (the improvement comprising)
        "In a data processing system comprising a processor and a memory, "
        "the improvement comprising a cache memory coupled to the processor "
        "for storing frequently accessed data.",
        # Claim 3: Markush open-transition (selected from the group comprising)
        "A composition of matter comprising a polymer selected from the group "
        "comprising polyethylene, polypropylene, and polystyrene.",
        # Claim 4: Omnibus (short + omnibus language)
        "An apparatus substantially as shown in the drawings.",
        # Claim 5: Wherein-comma violation ("wherein in response" — "in" requires comma after wherein)
        "The system of claim 2, wherein in response to receiving a signal, "
        "the cache memory stores the data.",
        # Claim 6: Extra periods (".." mid-claim)
        "The system of claim 2, wherein the cache memory has a size of 256 KB.. "
        "and an access time of 10 ns.",
    ]
    for text in claims:
        para = doc.add_paragraph(text)
        _set_para_num(para, "2")

    doc.add_paragraph("ABSTRACT")
    doc.add_paragraph(
        "A testbed document for verifying US claim-defect checks. The document recites six claims, "
        "each engineered to exercise a specific defect class in isolation."
    )

    return _doc_to_bytes(doc)


def _build_us_cluster2_spec_defects() -> bytes:
    """Engineered US patent exercising 4 zero-coverage spec-level defects.

    Phase E cluster 2: tracked changes (Word revisions), non-sequential
    paragraph numbering, missing required section, figure-xref orphan.
    """
    doc = Document()
    _add_us_numbering(doc)
    # Third numbering definition with startAt=5 to create a gap in
    # paragraph_numberings (e.g., [1, 2, 3, 5, 6]).
    _add_us_numbering_gap_id(doc, abstract_id="2", num_id="3", start_at=5)

    headers = {
        "TITLE OF THE INVENTION",
        "Widget Management System",
        "FIELD OF THE INVENTION",
        "SUMMARY OF THE INVENTION",
        "BRIEF DESCRIPTION OF THE DRAWINGS",
        "DETAILED DESCRIPTION OF THE PREFERRED EMBODIMENTS",
    }

    # Pre-gap spec paragraphs — numId=1 counts 1, 2, 3. Deliberately omits
    # BACKGROUND OF THE INVENTION to trigger the missing-required-section
    # check. Brief Description of Drawings mentions FIG. 1 and FIG. 2 to
    # set up the figure-xref orphan (only FIG. 1 will appear in the
    # Detailed Description below).
    pre_gap = [
        ("TITLE OF THE INVENTION", None),
        ("Widget Management System", None),
        ("FIELD OF THE INVENTION", None),
        ("The present invention relates to a widget management system.", "1"),
        ("SUMMARY OF THE INVENTION", None),
        (
            "A widget management system includes a controller and a plurality of widgets "
            "operatively coupled to the controller.",
            "1",
        ),
        ("BRIEF DESCRIPTION OF THE DRAWINGS", None),
        (
            "FIG. 1 is a schematic overview of the widget management system. "
            "FIG. 2 is a detail view of a single widget.",
            "1",
        ),
    ]
    for text, num in pre_gap:
        para = doc.add_paragraph(text)
        if num and text not in headers:
            _set_para_num(para, num)

    # Post-gap spec paragraphs — numId=3 starts at 5, producing a gap from
    # 3 to 5. The detailed description only references FIG. 1, leaving
    # FIG. 2 orphaned vs. the Brief Description.
    doc.add_paragraph("DETAILED DESCRIPTION OF THE PREFERRED EMBODIMENTS")
    para_detailed_1 = doc.add_paragraph(
        "Referring to FIG. 1, the widget management system 100 includes a controller 10 and "
        "a plurality of widgets 20 operatively coupled to the controller."
    )
    _set_para_num(para_detailed_1, "3")
    # Inject a tracked-change (w:ins) element on this paragraph so
    # detect_tracked_changes fires.
    _inject_tracked_insertion(para_detailed_1, "pending revision")

    para_detailed_2 = doc.add_paragraph(
        "The controller 10 sends command signals to each widget 20 over a communication bus."
    )
    _set_para_num(para_detailed_2, "3")

    doc.add_paragraph("CLAIMS")
    claim_para = doc.add_paragraph(
        "A widget management system comprising: a controller; and a plurality of widgets "
        "operatively coupled to the controller."
    )
    _set_para_num(claim_para, "2")

    doc.add_paragraph("ABSTRACT")
    doc.add_paragraph(
        "A widget management system includes a controller and a plurality of widgets "
        "operatively coupled to the controller. The controller sends command signals to "
        "each widget over a communication bus, enabling coordinated operation across the "
        "fleet of widgets."
    )

    return _doc_to_bytes(doc)


def _build_us_cluster3_single_figure() -> bytes:
    """Engineered US patent with a single figure mislabeled as FIG. 1.

    Phase E cluster 3: single-figure patents must refer to "The Figure"
    rather than "FIG. 1" per MPEP convention. Only FIG. 1 is mentioned
    anywhere in the document.
    """
    doc = Document()
    _add_us_numbering(doc)

    headers = {
        "TITLE OF THE INVENTION",
        "Single-Figure Apparatus",
        "FIELD OF THE INVENTION",
        "BACKGROUND OF THE INVENTION",
        "SUMMARY OF THE INVENTION",
        "BRIEF DESCRIPTION OF THE DRAWINGS",
        "DETAILED DESCRIPTION OF THE PREFERRED EMBODIMENTS",
    }
    spec = [
        "TITLE OF THE INVENTION",
        "Single-Figure Apparatus",
        "FIELD OF THE INVENTION",
        "The present invention relates to a single-figure apparatus.",
        "BACKGROUND OF THE INVENTION",
        "Conventional apparatus designs have limitations.",
        "SUMMARY OF THE INVENTION",
        "An apparatus includes a base and a cover coupled to the base.",
        "BRIEF DESCRIPTION OF THE DRAWINGS",
        "FIG. 1 is a perspective view of the apparatus.",
        "DETAILED DESCRIPTION OF THE PREFERRED EMBODIMENTS",
        "Referring to FIG. 1, the apparatus 100 includes a base 10 and a cover 20 coupled to the base.",
    ]
    for text in spec:
        para = doc.add_paragraph(text)
        if text not in headers:
            _set_para_num(para, "1")

    doc.add_paragraph("CLAIMS")
    claim_para = doc.add_paragraph(
        "An apparatus comprising: a base; and a cover coupled to the base."
    )
    _set_para_num(claim_para, "2")

    doc.add_paragraph("ABSTRACT")
    doc.add_paragraph(
        "An apparatus includes a base and a cover coupled to the base, providing a compact "
        "single-piece enclosure suitable for a range of applications."
    )

    return _doc_to_bytes(doc)


def _build_cn_amend_triggers() -> bytes:
    """Engineered CN patent triggering 2 AMEND-status checks.

    Phase E CN AMEND-triggers: commercial language in abstract (最佳) and
    abstract over 300 characters. Section ordering is not exercised here —
    see the note in test_cn_specification.py::TestSectionOrdering for the
    architectural limitation (CnPatentDocument stores sections by named
    field, so the check can never fail at the pipeline layer).
    """
    doc = Document()

    # Abstract: >300 chars AND contains commercial superlative "最佳"
    # ("best"). Character count excludes whitespace/newlines per
    # sections_cn.py:589.
    abstract_text = (
        "本发明提供一种最佳的数据处理方法，包括数据预处理步骤和特征提取步骤。"
        "数据预处理步骤用于对原始数据进行清洗和归一化处理，消除噪声和异常值。"
        "特征提取步骤采用深度学习模型从清洗后的数据中提取高维特征向量。"
        "该深度学习模型具有多层卷积结构，能够有效提取数据的空间特征和时序特征。"
        "本方法在图像识别、自然语言处理、语音识别、医疗诊断、金融风控等诸多领域展现出卓越的应用效果。"
        "系统能够有效处理大规模高维数据集，实现快速准确的数据分类和预测。"
        "本方法适用于多种复杂场景，显著提升系统处理效率和分类准确率。"
        "广泛的应用前景和重要的实用价值，使其成为相关领域的关键技术突破。"
        "此外，本方法还具备良好的可扩展性和稳定性，适合在大规模生产环境中长期稳定运行。"
        "综上所述，本发明通过创新的技术架构为数据处理领域提供了高效可靠的解决方案。"
    )
    doc.add_paragraph(abstract_text)
    _add_cn_section_break(doc, "摘要附图")

    _add_cn_section_break(doc, "权利要求书")
    doc.add_paragraph("1. 一种数据处理方法，包括：")
    doc.add_paragraph("数据预处理步骤，用于对原始数据进行清洗；及")
    doc.add_paragraph("特征提取步骤，用于从清洗后的数据中提取特征向量；")
    doc.add_paragraph("其特征在于，所述特征提取步骤采用深度学习模型。")
    _add_cn_section_break(doc, "说明书")

    doc.add_paragraph("一种数据处理方法")
    doc.add_paragraph("技术领域")
    doc.add_paragraph("本发明涉及一种数据处理方法。")
    doc.add_paragraph("背景技术")
    doc.add_paragraph("现有数据分类方法效率较低。")
    doc.add_paragraph("发明内容")
    doc.add_paragraph("本发明提供一种数据处理方法。")
    doc.add_paragraph("附图说明")
    doc.add_paragraph("图1是本发明的流程图。")
    doc.add_paragraph("具体实施方式")
    doc.add_paragraph("下面结合实施例对本发明进行详细说明。")
    _add_cn_section_break(doc, "说明书附图")

    first_section = doc.sections[0]
    first_section.header.is_linked_to_previous = False
    hp = first_section.header.paragraphs[0]
    hp.text = "说明书摘要"

    return _doc_to_bytes(doc)


def _build_cn_section_ordering_violation() -> bytes:
    """Engineered CN patent with spec subsection headers out of canonical order.

    Phase 9 #66 trigger: places 具体实施方式 before 发明内容 in the spec body.
    Section content itself is otherwise valid; the violation is purely
    ordering. Mirrors the MPEP-ordered-spec-reused-for-CNIPA drafter error.
    """
    doc = Document()

    # Abstract (passes charCount, no commercial language)
    doc.add_paragraph(
        "本发明提供一种数据处理方法，包括数据预处理步骤和特征提取步骤。"
        "数据预处理步骤用于对原始数据进行清洗。"
        "特征提取步骤用于提取高维特征向量。"
    )
    _add_cn_section_break(doc, "摘要附图")

    _add_cn_section_break(doc, "权利要求书")
    doc.add_paragraph("1. 一种数据处理方法，其特征在于，包括数据预处理步骤。")
    _add_cn_section_break(doc, "说明书")

    # Title + out-of-order subsection headers:
    #   技术领域 (0) → 具体实施方式 (4) → 发明内容 (2) → 附图说明 (3)
    # Indices 0, 4, 2, 3 are NOT strictly increasing.
    doc.add_paragraph("一种数据处理方法")
    doc.add_paragraph("技术领域")
    doc.add_paragraph("本发明涉及一种数据处理方法。")
    doc.add_paragraph("具体实施方式")
    doc.add_paragraph("下面结合实施例对本发明进行详细说明。")
    doc.add_paragraph("发明内容")
    doc.add_paragraph("本发明提供一种数据处理方法。")
    doc.add_paragraph("附图说明")
    doc.add_paragraph("图1是本发明的流程图。")
    _add_cn_section_break(doc, "说明书附图")

    first_section = doc.sections[0]
    first_section.header.is_linked_to_previous = False
    hp = first_section.header.paragraphs[0]
    hp.text = "说明书摘要"

    return _doc_to_bytes(doc)


def _build_tw_minimal(claims_text: list[str], symbol_lines: list[str] | None = None,
                      embodiment_lines: list[str] | None = None) -> bytes:
    """Build a minimal but complete TW .docx for targeted tests."""
    doc = Document()
    lines = [
        "【發明名稱】",
        "散熱裝置",
        "【技術領域】",
        "【0001】本發明係關於一種散熱裝置。",
        "【先前技術】",
        "【0002】習知技術存在問題。",
        "【發明內容】",
        "【0003】本發明提供一種散熱裝置。",
        "【圖式簡單說明】",
        "【0004】第1圖係散熱裝置之示意圖。",
        "【實施方式】",
    ]
    if embodiment_lines:
        lines.extend(embodiment_lines)
    else:
        lines.append("【0005】請參閱第1圖，散熱裝置100包括一基座10。")
    if symbol_lines is not None:
        lines.append("【符號說明】")
        lines.extend(symbol_lines)
    else:
        lines.append("【符號說明】")
        lines.extend(["100  散熱裝置", "10   基座"])
    lines.append("【申請專利範圍】")
    lines.extend(claims_text)
    lines.extend([
        "【摘要】",
        "本發明提供一種散熱裝置。",
        "【代表圖】",
        "第1圖",
    ])
    for line in lines:
        doc.add_paragraph(line)
    return _doc_to_bytes(doc)


def _build_tw_section_ordering_violation() -> bytes:
    """Engineered TW patent with bracket-header sections out of canonical order.

    Phase 9 #66 trigger: places 【實施方式】 before 【發明內容】 in the spec
    body. Section content is otherwise valid; violation is purely ordering.
    """
    doc = Document()
    lines = [
        "【發明名稱】",
        "散熱裝置",
        "【技術領域】",
        "【0001】本發明係關於一種散熱裝置。",
        "【先前技術】",
        "【0002】習知技術存在問題。",
        # 實施方式 (idx 4) BEFORE 發明內容 (idx 2) — violation
        "【實施方式】",
        "【0003】請參閱第1圖，散熱裝置100包括一基座10。",
        "【發明內容】",
        "【0004】本發明提供一種散熱裝置。",
        "【圖式簡單說明】",
        "【0005】第1圖係散熱裝置之示意圖。",
        "【符號說明】",
        "100  散熱裝置",
        "10   基座",
        "【申請專利範圍】",
        "1. 一種散熱裝置，其特徵在於包括一基座。",
        "【摘要】",
        "本發明提供一種散熱裝置。",
        "【代表圖】",
        "第1圖",
    ]
    for line in lines:
        doc.add_paragraph(line)
    return _doc_to_bytes(doc)


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


class TestTwInventionAllPass:
    def test_no_amend_findings(self):
        result = analyze_bytes(
            _build_tw_invention_all_pass(), "tw_invention.docx", Jurisdiction.TW
        )
        assert result.jurisdiction == Jurisdiction.TW
        assert result.likely_patent is True

        all_checks = (
            result.tw_specification_checks
            + result.tw_claims_checks
            + result.tw_abstract_checks
            + result.tw_drawings_checks
        )
        assert len(all_checks) > 0
        amend_items = [c for c in all_checks if c.status == "amend"]
        assert amend_items == [], f"Unexpected AMENDs: {[c.message_key for c in amend_items]}"

    def test_check_sections_populated(self):
        result = analyze_bytes(
            _build_tw_invention_all_pass(), "tw_invention.docx", Jurisdiction.TW
        )
        assert len(result.tw_specification_checks) > 0
        assert len(result.tw_claims_checks) > 0
        assert len(result.tw_abstract_checks) > 0

    def test_claims_parsed(self):
        result = analyze_bytes(
            _build_tw_invention_all_pass(), "tw_invention.docx", Jurisdiction.TW
        )
        # Claim 5 is 引用記載型式 (quoted-reference independent) — preamble
        # `一種散熱系統` declares a new subject; the body's `如請求項1所述
        # 之散熱裝置` is incorporation-by-reference, not dependency.
        # Claims 2/3/4 are true dependents.
        assert result.independent_claims_count == 2
        assert result.dependent_claims_count == 3


class TestTwUtilityModelAllPass:
    def test_utility_model_detected(self):
        result = analyze_bytes(
            _build_tw_utility_model_all_pass(), "tw_utility.docx", Jurisdiction.TW
        )
        assert result.jurisdiction == Jurisdiction.TW
        assert result.likely_patent is True

        # Check patent_type via the parsed document — we verify by checking
        # that 其改良在於 claims are accepted (no transition_phrase AMEND)
        all_checks = (
            result.tw_specification_checks
            + result.tw_claims_checks
            + result.tw_abstract_checks
            + result.tw_drawings_checks
        )
        amend_items = [c for c in all_checks if c.status == "amend"]
        assert amend_items == [], f"Unexpected AMENDs: {[c.message_key for c in amend_items]}"

    def test_claims_parsed(self):
        result = analyze_bytes(
            _build_tw_utility_model_all_pass(), "tw_utility.docx", Jurisdiction.TW
        )
        assert result.independent_claims_count == 1
        assert result.dependent_claims_count == 2


class TestTwInventionMultiFail:
    def test_multiple_findings(self):
        result = analyze_bytes(
            _build_tw_invention_multi_fail(), "tw_fail.docx", Jurisdiction.TW
        )
        assert result.likely_patent is True

        all_checks = (
            result.tw_specification_checks
            + result.tw_claims_checks
            + result.tw_abstract_checks
        )
        non_pass = [c for c in all_checks if c.status in ("amend", "verify")]
        assert len(non_pass) >= 5, (
            f"Expected >=5 findings, got {len(non_pass)}: "
            f"{[c.message_key for c in non_pass]}"
        )

    def test_specific_checks_triggered(self):
        result = analyze_bytes(
            _build_tw_invention_multi_fail(), "tw_fail.docx", Jurisdiction.TW
        )
        all_checks = (
            result.tw_specification_checks
            + result.tw_claims_checks
            + result.tw_abstract_checks
        )
        keys = {c.message_key for c in all_checks if c.status in ("amend", "verify")}

        expected_keys = {
            "check.tw.spec.symbolTablePresence.amend",
            "check.tw.spec.paragraphNumbering.amendGap",
            "check.tw.claims.forwardDependency.amend",
            "check.tw.abstract.commercialLanguage.amend",
        }
        for key in expected_keys:
            assert key in keys, f"Expected {key} in findings, got: {keys}"


class TestCnDocxAllPass:
    def test_no_amend_findings(self):
        result = analyze_bytes(
            _build_cn_docx_all_pass(), "cn_patent.docx", Jurisdiction.CN
        )
        assert result.jurisdiction == Jurisdiction.CN
        assert result.likely_patent is True

        all_checks = (
            result.cn_specification_checks
            + result.cn_claims_checks
            + result.cn_abstract_checks
            + result.cn_drawings_checks
        )
        assert len(all_checks) > 0
        amend_items = [c for c in all_checks if c.status == "amend"]
        assert amend_items == [], f"Unexpected AMENDs: {[c.message_key for c in amend_items]}"

    def test_sections_populated(self):
        result = analyze_bytes(
            _build_cn_docx_all_pass(), "cn_patent.docx", Jurisdiction.CN
        )
        assert len(result.cn_specification_checks) > 0
        assert len(result.cn_claims_checks) > 0
        assert len(result.cn_abstract_checks) > 0


class TestCnDocxMultiFail:
    def test_multiple_findings(self):
        result = analyze_bytes(
            _build_cn_docx_multi_fail(), "cn_fail.docx", Jurisdiction.CN
        )
        all_checks = (
            result.cn_specification_checks
            + result.cn_claims_checks
            + result.cn_abstract_checks
        )
        non_pass = [c for c in all_checks if c.status in ("amend", "verify")]
        assert len(non_pass) >= 4, (
            f"Expected >=4 findings, got {len(non_pass)}: "
            f"{[c.message_key for c in non_pass]}"
        )

    def test_specific_checks_triggered(self):
        result = analyze_bytes(
            _build_cn_docx_multi_fail(), "cn_fail.docx", Jurisdiction.CN
        )
        all_checks = (
            result.cn_specification_checks
            + result.cn_claims_checks
            + result.cn_abstract_checks
        )
        keys = {c.message_key for c in all_checks if c.status in ("amend", "verify")}

        assert "check.cn.spec.requiredSections.amend" in keys, f"Missing requiredSections in {keys}"
        assert "check.cn.claims.sequential.amend" in keys, f"Missing sequential in {keys}"


class TestUsFullLength:
    def test_no_amend_findings(self):
        result = analyze_bytes(
            _build_us_full_length(), "us_patent.docx", Jurisdiction.US
        )
        assert result.jurisdiction == Jurisdiction.US
        assert result.likely_patent is True

        # US uses to_report_data() for check aggregation
        report = result.to_report_data()
        all_checks = (
            report.specification_checks
            + report.claims_checks
            + report.abstract_checks
            + report.drawings_checks
        )
        amend_items = [c for c in all_checks if c.status == "amend"]
        assert amend_items == [], f"Unexpected AMENDs: {[c.message_key for c in amend_items]}"

    def test_claims_parsed(self):
        result = analyze_bytes(
            _build_us_full_length(), "us_patent.docx", Jurisdiction.US
        )
        assert result.independent_claims_count == 3
        assert result.dependent_claims_count == 7

    def test_sections_all_present(self):
        report = analyze_bytes(
            _build_us_full_length(), "us_patent.docx", Jurisdiction.US
        ).to_report_data()
        assert len(report.specification_checks) > 0
        assert len(report.claims_checks) > 0
        assert len(report.abstract_checks) > 0
        assert len(report.drawings_checks) > 0


class TestUsCluster1Defects:
    """Phase E cluster 1: one testspec exercises 6 zero-coverage defect checks."""

    TARGET_KEYS = {
        "claims.crmNonTransitory",
        "claims.jepsonPriorArt",
        "claims.markushOpenTransition",
        "claims.omnibusClaim",
        "claims.whereinComma",
        "claims.extraPeriod",
    }

    def _emitted_keys(self) -> set[str]:
        result = analyze_bytes(
            _build_us_cluster1_defects(), "us_cluster1.docx", Jurisdiction.US
        )
        report = result.to_report_data()
        return {c.message_key for c in report.all_checks if c.message_key}

    def test_all_six_target_keys_emitted(self):
        keys = self._emitted_keys()
        missing = self.TARGET_KEYS - keys
        assert not missing, f"Missing defect-check emissions: {sorted(missing)}"

    def test_crm_non_transitory_detected(self):
        assert "claims.crmNonTransitory" in self._emitted_keys()

    def test_jepson_detected(self):
        assert "claims.jepsonPriorArt" in self._emitted_keys()

    def test_markush_open_transition_detected(self):
        assert "claims.markushOpenTransition" in self._emitted_keys()

    def test_omnibus_detected(self):
        assert "claims.omnibusClaim" in self._emitted_keys()

    def test_wherein_comma_detected(self):
        assert "claims.whereinComma" in self._emitted_keys()

    def test_extra_period_detected(self):
        assert "claims.extraPeriod" in self._emitted_keys()


class TestUsCluster2SpecDefects:
    """Phase E cluster 2: 4 zero-coverage spec-level defect checks."""

    TARGET_KEYS = {
        "check.spec.trackedChanges.amend",
        "check.spec.paragraphSequential.amend",
        "checks.required_sections_missing",
        "checks.figure_xref_orphaned_brief",
    }

    def _emitted_keys(self) -> set[str]:
        result = analyze_bytes(
            _build_us_cluster2_spec_defects(), "us_cluster2.docx", Jurisdiction.US
        )
        report = result.to_report_data()
        return {c.message_key for c in report.all_checks if c.message_key}

    def test_all_four_target_keys_emitted(self):
        keys = self._emitted_keys()
        missing = self.TARGET_KEYS - keys
        assert not missing, f"Missing defect-check emissions: {sorted(missing)}"

    def test_tracked_changes_detected(self):
        assert "check.spec.trackedChanges.amend" in self._emitted_keys()

    def test_non_sequential_paragraphs_detected(self):
        assert "check.spec.paragraphSequential.amend" in self._emitted_keys()

    def test_missing_required_sections_detected(self):
        assert "checks.required_sections_missing" in self._emitted_keys()

    def test_figure_xref_orphaned_detected(self):
        assert "checks.figure_xref_orphaned_brief" in self._emitted_keys()


class TestUsCluster3SingleFigure:
    """Phase E cluster 3: single-figure patent mislabeled as 'FIG. 1'."""

    def _emitted_keys(self) -> set[str]:
        result = analyze_bytes(
            _build_us_cluster3_single_figure(), "us_cluster3.docx", Jurisdiction.US
        )
        report = result.to_report_data()
        return {c.message_key for c in report.all_checks if c.message_key}

    def test_single_figure_wrong_label_detected(self):
        assert "check.drawings.singleFigure.amend" in self._emitted_keys()


class TestCnAmendTriggers:
    """Phase E CN cluster: AMEND-status abstract defect checks that can be
    triggered via the full pipeline."""

    TARGET_KEYS = {
        "check.cn.abstract.commercialLanguage.amend",
        "check.cn.abstract.charCount.amend",
    }

    def _emitted_keys(self) -> set[str]:
        result = analyze_bytes(
            _build_cn_amend_triggers(), "cn_amend.docx", Jurisdiction.CN
        )
        all_checks = (
            result.cn_specification_checks
            + result.cn_claims_checks
            + result.cn_abstract_checks
            + result.cn_drawings_checks
        )
        return {c.message_key for c in all_checks if c.message_key}

    def test_all_target_keys_emitted(self):
        keys = self._emitted_keys()
        missing = self.TARGET_KEYS - keys
        assert not missing, f"Missing defect-check emissions: {sorted(missing)}"

    def test_commercial_language_detected(self):
        assert "check.cn.abstract.commercialLanguage.amend" in self._emitted_keys()

    def test_abstract_over_300_chars_detected(self):
        assert "check.cn.abstract.charCount.amend" in self._emitted_keys()


class TestCnSectionOrderingViolation:
    """Phase 9 #66 CN: spec subsection headers emitted out of canonical order
    trigger ``check.cn.spec.sectionOrdering.amend`` via the full pipeline."""

    def _emitted_keys(self) -> set[str]:
        result = analyze_bytes(
            _build_cn_section_ordering_violation(),
            "cn_section_order.docx",
            Jurisdiction.CN,
        )
        all_checks = (
            result.cn_specification_checks
            + result.cn_claims_checks
            + result.cn_abstract_checks
            + result.cn_drawings_checks
        )
        return {c.message_key for c in all_checks if c.message_key}

    def test_section_ordering_amend_detected(self):
        keys = self._emitted_keys()
        assert "check.cn.spec.sectionOrdering.amend" in keys, (
            f"Expected section ordering AMEND; got {sorted(keys)}"
        )

    def test_section_ordering_pass_not_emitted(self):
        keys = self._emitted_keys()
        assert "check.cn.spec.sectionOrdering.pass" not in keys


class TestTwSectionOrderingViolation:
    """Phase 9 #66 TW: 【】bracket headers emitted out of canonical order
    trigger ``check.tw.spec.sectionOrdering.amend`` via the full pipeline."""

    def _emitted_keys(self) -> set[str]:
        result = analyze_bytes(
            _build_tw_section_ordering_violation(),
            "tw_section_order.docx",
            Jurisdiction.TW,
        )
        all_checks = (
            result.tw_specification_checks
            + result.tw_claims_checks
            + result.tw_abstract_checks
            + result.tw_drawings_checks
        )
        return {c.message_key for c in all_checks if c.message_key}

    def test_section_ordering_amend_detected(self):
        keys = self._emitted_keys()
        assert "check.tw.spec.sectionOrdering.amend" in keys, (
            f"Expected section ordering AMEND; got {sorted(keys)}"
        )

    def test_section_ordering_pass_not_emitted(self):
        keys = self._emitted_keys()
        assert "check.tw.spec.sectionOrdering.pass" not in keys


class TestCrossJurisdictionMismatch:
    def test_tw_doc_analyzed_as_cn(self):
        """TW 【】bracket headers should not be recognized by CN parser."""
        result = analyze_bytes(
            _build_tw_invention_all_pass(), "tw_as_cn.docx", Jurisdiction.CN
        )
        # CN parser should find no meaningful sections from TW brackets.
        # The document is still parsed but sections will be mostly empty.
        # Check that the pipeline doesn't crash.
        assert result.jurisdiction == Jurisdiction.CN

    def test_cn_doc_analyzed_as_tw(self):
        """CN section headers (no 【】brackets) should not be recognized by TW parser."""
        result = analyze_bytes(
            _build_cn_docx_all_pass(), "cn_as_tw.docx", Jurisdiction.TW
        )
        # TW parser should find no bracket headers in CN document
        assert result.jurisdiction == Jurisdiction.TW


class TestTwAntecedentBasisChain:
    def test_antecedent_basis_chain(self):
        claims = [
            "1. 一種散熱裝置，包括：",
            "一基座(10)；及",
            "一散熱鰭片(20)，設置於該基座(10)上；",
            "其特徵在於，該基座(10)具有一導熱層(30)。",
            "2. 如請求項1所述之散熱裝置，其中該導熱層(30)包括一銅箔層(31)。",
            "3. 如請求項2所述之散熱裝置，其中該銅箔層(31)之厚度為0.1至0.5毫米。",
        ]
        data = _build_tw_minimal(claims)
        result = analyze_bytes(data, "tw_antecedent.docx", Jurisdiction.TW)

        # Find antecedent basis check
        ab_checks = [
            c for c in result.tw_claims_checks
            if "antecedentBasis" in c.message_key
        ]
        assert len(ab_checks) > 0, "antecedentBasis check not found"
        # Should be PASS or VERIFY (not AMEND) — all 該X have matching 一X introductions
        for check in ab_checks:
            assert check.status != "amend", (
                f"antecedentBasis should not be AMEND: {check.message_key} / {check.details}"
            )


class TestTwSymbolTableRangeNumerals:
    def test_range_numerals_parsed(self):
        symbol_lines = [
            "100  散熱裝置",
            "S21~S25  感測器組",
            "3001~3010  連接元件",
            "10   基座",
        ]
        embodiment_lines = [
            "【0005】請參閱第1圖，散熱裝置100包括一基座10、感測器組S21~S25、及連接元件3001~3010。",
        ]
        data = _build_tw_minimal(
            claims_text=[
                "1. 一種散熱裝置，包括：",
                "一基座(10)；",
                "其特徵在於，該基座(10)具有一導熱層。",
            ],
            symbol_lines=symbol_lines,
            embodiment_lines=embodiment_lines,
        )
        result = analyze_bytes(data, "tw_symbol_range.docx", Jurisdiction.TW)

        # Pipeline should not crash on range numerals
        assert result.jurisdiction == Jurisdiction.TW
        assert result.likely_patent is True

        # Check that symbolTableConsistency is present
        consistency_checks = [
            c for c in result.tw_specification_checks
            if "symbolTable" in c.message_key and "Consistency" in c.message_key
        ]
        assert len(consistency_checks) > 0, "symbolTableConsistency check not found"


# ---------------------------------------------------------------------------
# Real TW patent fixture tests (firm-variant format)
# ---------------------------------------------------------------------------

class TestTwRealFixtures:
    """Integration tests using real TW patent .docx files from a Taiwan patent firm.

    These files use firm-variant headers (【中文發明名稱】, 【發明申請專利範圍】, etc.)
    and Word numbering (w:numPr) for claims.
    """

    REAL_FIXTURE_DIR = Path(__file__).parent / "fixtures" / "tw"
    REAL_PATTERNS = sorted(REAL_FIXTURE_DIR.glob("*派譯版*"))

    @pytest.fixture(params=[p.name for p in REAL_PATTERNS], ids=[p.stem[:30] for p in REAL_PATTERNS])
    def tw_result(self, request):
        fpath = self.REAL_FIXTURE_DIR / request.param
        with open(fpath, "rb") as f:
            data = f.read()
        return analyze_bytes(data, request.param, Jurisdiction.TW)

    def test_is_patent(self, tw_result):
        assert tw_result.likely_patent is True

    def test_claims_detected(self, tw_result):
        assert len(tw_result.claims) > 0, "No claims detected"
        assert tw_result.independent_claims_count > 0

    def test_abstract_populated(self, tw_result):
        assert tw_result.abstract_word_count > 0, "Abstract char count is 0"

    def test_title_not_falsely_flagged(self, tw_result):
        title_checks = [
            c for c in tw_result.tw_specification_checks
            if "title" in c.message_key
        ]
        for check in title_checks:
            if check.status == "amend":
                assert "prohibited" not in (check.message or "").lower(), (
                    f"False positive title check: {check.message}"
                )


class TestTwRealFixtureUtilityModel:
    """Verify patent type detection for the utility model fixture."""

    def test_utility_model_detection(self):
        fpath = Path(__file__).parent / "fixtures" / "tw" / "110P000840US.JP.DE派譯版-FV.DOCX"
        if not fpath.exists():
            pytest.skip("Utility model fixture not available")
        with open(fpath, "rb") as f:
            data = f.read()
        result = analyze_bytes(data, fpath.name, Jurisdiction.TW)
        # Check patent type via specification checks — utility model terms
        # should not trigger type terminology mismatch
        type_checks = [
            c for c in result.tw_specification_checks
            if "patentTypeTerminology" in c.message_key
        ]
        assert any(c.status == "pass" for c in type_checks), (
            "Utility model patent type terminology check should pass"
        )

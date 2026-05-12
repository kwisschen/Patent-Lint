# SPDX-License-Identifier: LicenseRef-PolyForm-Strict-1.0.0
# Copyright (c) 2025–2026 Christopher Chen
"""Scaffolding-stage assertions for the EPC pipeline.

These tests do NOT exercise any real EPC check logic. They lock in that the
EPC routing path through ``analyze_bytes`` works end-to-end without crashing,
and that the empty-stub pipeline returns a well-formed ``AnalysisResult`` /
``ReportData`` pair so the frontend can render an EPC analysis page (empty
sections, no findings) the moment the picker offers EPC.

When real EPC checks ship per the implementation plan, these tests stay as
the regression gate — every G-group commit should leave them passing.
"""

from __future__ import annotations

import io

import pytest
from docx import Document

from patentlint.models import AnalysisResult, Jurisdiction, ReportData
from patentlint.pipeline import _run_epc_pipeline, analyze_bytes


def _make_minimal_docx_bytes() -> bytes:
    """Build a one-paragraph English .docx in memory."""
    doc = Document()
    doc.add_paragraph("Example title")
    doc.add_paragraph("This is a placeholder English paragraph for EPC pipeline scaffolding tests.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_run_epc_pipeline_stub_returns_well_formed_result():
    """Pipeline runs G1 + G2 spec checks (9) + G3 (4) + G4-G6 claims (17) + G7 (4)."""
    result = _run_epc_pipeline("any english text", suggested_jurisdiction=None)
    assert isinstance(result, AnalysisResult)
    assert result.jurisdiction == Jurisdiction.EPC
    assert len(result.epc_specification_checks) == 9
    assert len(result.epc_drawings_checks) == 4
    assert len(result.epc_claims_checks) == 17
    assert len(result.epc_abstract_checks) == 4


def test_analyze_bytes_routes_epc_jurisdiction_without_crashing():
    content = _make_minimal_docx_bytes()
    result = analyze_bytes(content, "epc_sample.docx", jurisdiction=Jurisdiction.EPC)
    assert isinstance(result, AnalysisResult)
    assert result.jurisdiction == Jurisdiction.EPC


def test_analyze_bytes_rejects_non_docx_for_epc():
    with pytest.raises(ValueError, match="Unsupported file type for EPC"):
        analyze_bytes(b"", "epc_sample.xml", jurisdiction=Jurisdiction.EPC)


def test_epc_report_data_adapter_round_trips():
    result = _run_epc_pipeline("any english text")
    report = result.to_report_data()
    assert isinstance(report, ReportData)
    assert report.jurisdiction == Jurisdiction.EPC
    # G1 + G2 ship 9 spec checks; G3 ships 4 drawings; G4-G6 ship 17
    # claims; G7 ships 4 abstract.
    assert len(report.specification_checks) == 9
    assert len(report.drawings_checks) == 4
    assert len(report.claims_checks) == 17
    assert len(report.abstract_checks) == 4


def test_epc_pipeline_attaches_rubric_grade():
    """EPC pipeline must attach a rubric_grade (Bug #1 regression).

    Pre-fix, _run_epc_pipeline created AnalysisResult without calling
    _attach_rubric_grade — so the hero dashed out (no letter, no score)
    even on a draft with complete sections.
    """
    result = _run_epc_pipeline("any english text")
    assert result.rubric_grade is not None
    # When all required sections are empty, the completeness gate fires
    # and emits a CompletenessGap instead of a letter; either outcome
    # means the grade was computed (not None due to missing call).
    has_grade_or_gap = (
        result.rubric_grade.completeness_gap is not None
        or result.rubric_grade.letter is not None
    )
    assert has_grade_or_gap


def test_epc_figure_xref_message_keys_rekey_to_epc_namespace():
    """check_figure_ref_consistency_epc must re-key US figure-xref keys.

    Pre-fix the .replace() looked for `checks.figureXref` (camelCase)
    but the US helper emits `checks.figure_xref_*` (snake_case), so the
    keys passed through unchanged and rendered the MPEP-citing en.json
    `checks.figure_xref_orphaned_detailed` text on EPC analyses.
    """
    from patentlint.analysis.epc_specification import check_figure_ref_consistency_epc

    # Mismatched figure references: brief mentions FIG. 1, FIG. 2, FIG. 3;
    # detailed only describes FIG. 1. FIG. 2 and FIG. 3 are orphaned-brief.
    full_text = (
        "Title\n"
        "BRIEF DESCRIPTION OF THE DRAWINGS\n"
        "FIG. 1 shows a thing. FIG. 2 shows another thing. FIG. 3 shows a third thing.\n"
        "DETAILED DESCRIPTION\n"
        "As shown in FIG. 1, the thing exists.\n"
    )
    items = check_figure_ref_consistency_epc(full_text)
    assert items, "expected at least one CheckItem for mismatched figure refs"
    for item in items:
        assert item.message_key.startswith("check.epc.spec.figureRefConsistency"), (
            f"message_key should be in EPC namespace, got: {item.message_key}"
        )
        # Inline message must not carry US MPEP citation.
        assert "MPEP" not in (item.message or "")


def test_epc_jurisdiction_mismatch_detects_us_draft():
    """EPC selection on a US-flavored draft should suggest US (Bug #4).

    Pre-fix the EPC→US detector required `epc_markers == 0` which
    failed on the common case of a US draft using `method of claim N`
    phrasing without explicitly mentioning USPTO / MPEP.
    """
    from patentlint.models import Jurisdiction
    from patentlint.parser.jurisdiction_mismatch import detect_jurisdiction_mismatch

    us_flavored = (
        "1. A method of doing a thing. 2. The method of claim 1, wherein "
        "the thing is improved. 3. The system of claim 1 further comprising "
        "a non-transitory computer-readable medium."
    )
    suggested = detect_jurisdiction_mismatch(us_flavored, Jurisdiction.EPC)
    assert suggested == "US"


def test_epc_detects_german_draft_as_unsupported():
    """German content in an EPC draft should flag as unsupported language.

    EPC v1 is English-only; the frontend renders an advisory banner
    when this field is set so DE drafters know the engine doesn't
    support their language yet.
    """
    from patentlint.parser.jurisdiction_mismatch import detect_epc_unsupported_language

    de_text = (
        "Die Erfindung betrifft eine Vorrichtung, die dadurch gekennzeichnet "
        "ist, dass sie ein erstes Bauteil und ein zweites Bauteil aufweist. "
        "Die Vorrichtung wird mit einer Steuereinheit verbunden und ist für "
        "die Ausführung eines Verfahrens auf der Grundlage gemessener Werte "
        "ausgelegt. Patentansprüche werden gemäß der Beschreibung formuliert."
    )
    assert detect_epc_unsupported_language(de_text) == "de"


def test_epc_detects_french_draft_as_unsupported():
    from patentlint.parser.jurisdiction_mismatch import detect_epc_unsupported_language

    fr_text = (
        "L'invention concerne un dispositif caractérisé en ce que le "
        "dispositif comprend un premier élément et un second élément. Le "
        "dispositif est connecté à une unité de commande qui est configurée "
        "pour exécuter un procédé selon les revendications. La description "
        "fournit un mode de réalisation détaillé de l'invention."
    )
    assert detect_epc_unsupported_language(fr_text) == "fr"


def test_epc_english_draft_not_flagged_as_unsupported():
    from patentlint.parser.jurisdiction_mismatch import detect_epc_unsupported_language

    en_text = (
        "The present invention relates to a device characterised in that "
        "the device comprises a first element and a second element. The "
        "device is connected to a control unit configured to execute a "
        "method according to the claims."
    )
    assert detect_epc_unsupported_language(en_text) is None

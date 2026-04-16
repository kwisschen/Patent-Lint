# SPDX-License-Identifier: PolyForm-Noncommercial-1.0.0
# Copyright (c) 2025 Christopher Chen
"""DOCX document loader using python-docx.

Extracts numbered paragraphs from patent .docx files, tracking paragraph and claim
numberings, missing endings, and restrictive wording — mirroring the Java
PatentAnalyzer's loadDocxContent() and processParagraph() logic.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from patentlint.analysis.specification import detect_restrictive_wording, has_valid_ending
from patentlint.parser.sections import extract_description_of_drawings_section

_INVISIBLE_CHARS = {
    "\u200b": "",   # ZWSP (zero-width space)
    "\u200c": "",   # ZWNJ (zero-width non-joiner)
    "\u200d": "",   # ZWJ (zero-width joiner)
    "\ufeff": "",   # BOM (byte order mark)
    "\u00a0": " ",  # NBSP (non-breaking space) -> regular space
}


def _normalize_unicode(text: str) -> str:
    """Strip invisible whitespace and replace NBSP with a regular space.

    Patent drafts pasted from PDFs or web sources frequently contain
    zero-width characters that silently break downstream string matching.
    Scope is strictly the 5 chars in ``_INVISIBLE_CHARS``; CJK punctuation
    and full-width characters are left untouched.
    """
    for src, dst in _INVISIBLE_CHARS.items():
        text = text.replace(src, dst)
    return text


@dataclass
class LoadedDocument:
    """Structured result of loading a patent .docx file.

    Contains the full document text plus numbering metadata extracted
    from Word's paragraph numbering XML.
    """

    full_text: str
    paragraph_numberings: list[int] = field(default_factory=list)
    claim_numberings: list[int] = field(default_factory=list)
    missing_ending_paragraphs: list[int] = field(default_factory=list)
    improper_spec_paragraphs: list[int] = field(default_factory=list)
    improper_spec_phrases: str = ""
    has_tracked_changes: bool = False


@dataclass
class DocxSection:
    """A Word document section with its page header text and body paragraphs.

    ``numpr_flags`` is a parallel list of booleans (same length as
    ``paragraphs``) flagging which paragraphs carry Word ``w:numPr``
    auto-numbering. It is populated by ``load_docx_cn`` (Phase 8c) so the
    downstream section-ID fallback chain can backfill synthetic ``"N. "``
    prefixes on claim paragraphs that lack a typed prefix. Older callers
    (and the minimal ``extract_cn_xml_from_zip`` path) can leave it empty.
    """

    header_text: str = ""
    paragraphs: list[str] = field(default_factory=list)
    numpr_flags: list[bool] = field(default_factory=list)


@dataclass
class LoadedTwDocument:
    """Structured result of loading a TW patent .docx file.

    ``paragraph_word_numbers`` is a parallel list aligned with ``paragraphs``.
    Each entry is the Word auto-numbering string (``"0001"``, ``"0109"``, …)
    Word would render for that paragraph, or ``None`` if the paragraph is
    not Word-auto-numbered. Populated by ``load_docx_tw`` so downstream
    checks can report flagged paragraphs using the same 【NNNN】 identifiers
    the drafter sees in Word rather than PatentLint-internal ordinals.
    """

    paragraphs: list[str] = field(default_factory=list)
    paragraph_word_numbers: list[str | None] = field(default_factory=list)
    has_tracked_changes: bool = False


@dataclass
class LoadedCnDocument:
    """Structured result of loading a CN patent .docx file."""

    sections: list[DocxSection] = field(default_factory=list)
    has_tracked_changes: bool = False


def detect_tracked_changes(doc) -> bool:
    """Check if a .docx document contains tracked changes (w:del or w:ins).

    Tracked changes mean the document has unresolved revisions that should
    be accepted or rejected before final filing.
    """
    body = doc.element.body
    return (
        len(body.findall('.//' + qn('w:del'))) > 0
        or len(body.findall('.//' + qn('w:ins'))) > 0
    )


def _get_paragraph_num_id(paragraph) -> str | None:
    """Extract the numbering ID from a paragraph's XML properties."""
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is None:
        return None
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None
    numId_elem = numPr.find(qn("w:numId"))
    if numId_elem is None:
        return None
    val = numId_elem.get(qn("w:val"))
    # numId 0 means "no numbering"
    if val is None or val == "0":
        return None
    return val


def _extract_numpr_claim_number(paragraph) -> int | None:
    """Probe whether a paragraph carries Word auto-numbering (``w:numPr``).

    Returns ``1`` as a sentinel when auto-numbering is present so the caller
    can backfill a synthetic ``"N. "`` prefix using a running counter
    (matching the TW/CN claim parser contract — both parsers rely on an
    Arabic-numeral + period prefix at paragraph start). Returns ``None``
    when the paragraph has no ``w:numPr`` or an explicit ``numId="0"``
    (which Word treats as "no numbering").

    Extracted from ``load_docx_tw`` at Phase 8c so ``load_docx_cn`` can
    share the same detection path.
    """
    num_id = _get_paragraph_num_id(paragraph)
    if num_id is None:
        return None
    # Sentinel value — the caller assigns the actual sequential claim
    # number. Word's numbering-start XML is not consulted because CN/TW
    # claim lists routinely restart per page-break section, which produces
    # misleading start values.
    return 1


def _get_paragraph_ilvl(paragraph) -> int:
    """Extract the indentation level from a paragraph's numbering properties."""
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is None:
        return 0
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return 0
    ilvl_elem = numPr.find(qn("w:ilvl"))
    if ilvl_elem is None:
        return 0
    val = ilvl_elem.get(qn("w:val"))
    return int(val) if val else 0


def _get_numbering_start(document, num_id: str, ilvl: int) -> int:
    """Get the start value for a numbering definition at a given level.

    Walks the numbering XML: num -> abstractNumId -> abstractNum -> lvl -> start.
    """
    numbering_part = document.part.numbering_part
    if numbering_part is None:
        return 1

    numbering_xml = numbering_part._element

    # Find the <w:num w:numId="..."> element
    for num_elem in numbering_xml.findall(qn("w:num")):
        if num_elem.get(qn("w:numId")) == num_id:
            abstract_ref = num_elem.find(qn("w:abstractNumId"))
            if abstract_ref is None:
                return 1
            abstract_num_id = abstract_ref.get(qn("w:val"))

            # Find the corresponding <w:abstractNum>
            for abstract_elem in numbering_xml.findall(qn("w:abstractNum")):
                if abstract_elem.get(qn("w:abstractNumId")) == abstract_num_id:
                    # Find the <w:lvl w:ilvl="..."> element
                    for lvl_elem in abstract_elem.findall(qn("w:lvl")):
                        if lvl_elem.get(qn("w:ilvl")) == str(ilvl):
                            start_elem = lvl_elem.find(qn("w:start"))
                            if start_elem is not None:
                                val = start_elem.get(qn("w:val"))
                                return int(val) if val else 1
                    return 1
            return 1
    return 1


_CLAIMS_HEADER = re.compile(r"^(CLAIMS|What is claimed is|I claim|We claim)$", re.IGNORECASE)


def load_docx(file_path: str | Path) -> LoadedDocument:
    """Load a patent .docx file and extract structured content.

    Mirrors the Java PatentAnalyzer's loadDocxContent() and processParagraph():
    - Detects Word numbering via XML (numId, abstractNum, lvl start values)
    - Tracks separate counters per numbering list ID
    - Splits paragraphs into spec (beforeClaims) and claims sections
    - Checks spec paragraphs for valid endings and restrictive wording
    - Detects description-of-drawings section for relaxed ending rules

    Args:
        file_path: Path to the .docx file.

    Returns:
        LoadedDocument with full text and numbering metadata.

    Raises:
        FileNotFoundError: If the file doesn't exist.
        ValueError: If the file is not a valid .docx.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    if path.suffix.lower() != ".docx":
        raise ValueError(f"Not a .docx file: {path}")

    try:
        doc = Document(str(path))
    except Exception as e:
        raise ValueError(f"Invalid .docx file: {e}") from e

    # Build full document text first (for section extraction)
    full_text_lines: list[str] = []
    for para in doc.paragraphs:
        full_text_lines.append(_normalize_unicode(para.text).strip())
    full_text = "\n".join(full_text_lines)

    # Extract drawings section for relaxed ending rules
    drawings_section = extract_description_of_drawings_section(full_text)

    # State tracking (mirrors Java's loadDocxContent)
    list_start_numbers: dict[str, int] = {}  # numId -> current counter
    before_claims = True
    paragraph_numberings: list[int] = []
    claim_numberings: list[int] = []
    missing_ending_paragraphs: list[int] = []
    improper_spec_paragraphs: list[int] = []
    improper_spec_phrases_parts: list[str] = []
    formatted_lines: list[str] = []

    for para in doc.paragraphs:
        paragraph_text = _normalize_unicode(para.text).strip()

        # Check for CLAIMS header
        if _CLAIMS_HEADER.match(paragraph_text):
            before_claims = False

        # Determine if in drawings section
        in_drawings = bool(drawings_section and paragraph_text and paragraph_text in drawings_section)

        # Extract numbering info
        num_id = _get_paragraph_num_id(para)

        if num_id is not None:
            ilvl = _get_paragraph_ilvl(para)
            start = _get_numbering_start(doc, num_id, ilvl)
            current_num = list_start_numbers.get(num_id, start)

            if paragraph_text:
                if before_claims:
                    # Check valid ending
                    if not has_valid_ending(paragraph_text, in_drawings):
                        missing_ending_paragraphs.append(current_num)

                    formatted_lines.append(f"[{current_num}] {paragraph_text}")
                    paragraph_numberings.append(current_num)

                    # Check restrictive wording
                    wording_result = detect_restrictive_wording(paragraph_text, current_num)
                    if wording_result.flagged_paragraphs:
                        for pn in wording_result.flagged_paragraphs:
                            if pn not in improper_spec_paragraphs:
                                improper_spec_paragraphs.append(pn)
                        improper_spec_phrases_parts.append(wording_result.formatted_phrases)
                else:
                    # Claims section
                    formatted_lines.append(f"{current_num}. {paragraph_text}")
                    claim_numberings.append(current_num)

            list_start_numbers[num_id] = current_num + 1
        else:
            # Non-numbered paragraph
            if paragraph_text:
                formatted_lines.append(paragraph_text)

    tracked_changes = detect_tracked_changes(doc)

    return LoadedDocument(
        full_text="\n".join(formatted_lines),
        paragraph_numberings=paragraph_numberings,
        claim_numberings=claim_numberings,
        missing_ending_paragraphs=missing_ending_paragraphs,
        improper_spec_paragraphs=improper_spec_paragraphs,
        improper_spec_phrases="".join(improper_spec_phrases_parts),
        has_tracked_changes=tracked_changes,
    )


def load_docx_tw(file_path: str | Path) -> LoadedTwDocument:
    """Load a TW patent .docx and return paragraph texts with claim numbering.

    TW .docx files use 【】bracket headers for sections and Word numbering
    (w:numPr) for claims. This function:
    - Extracts paragraph text as-is for section extraction
    - Detects claim paragraphs (w:numPr after a claims header) and prepends
      sequential ``N. `` so claims_tw.py can parse them
    - Detects tracked changes (w:del / w:ins)
    """
    path = Path(file_path)
    if not path.exists():
        msg = f"File not found: {path}"
        raise FileNotFoundError(msg)
    if path.suffix.lower() != ".docx":
        msg = f"Not a .docx file: {path}"
        raise ValueError(msg)

    try:
        doc = Document(str(path))
    except Exception as exc:
        msg = f"Failed to open .docx: {exc}"
        raise ValueError(msg) from exc

    # Claims section headers (same variants as sections_tw._SECTION_MAP)
    _TW_CLAIMS_HEADERS = {"申請專利範圍", "發明申請專利範圍", "新型申請專利範圍"}
    _TW_BRACKET_RE = re.compile(r"^【(.+?)】")

    paragraphs: list[str] = []
    paragraph_word_numbers: list[str | None] = []
    # Word auto-numbering runs one counter per numId. For TW patents the
    # 【NNNN】 paragraph-numbering format is typically a single numId applied
    # across the body; claim-list paragraphs use a separate numId and a
    # different format (e.g. N.) that we don't want to surface as the
    # body-paragraph 【NNNN】 label. Restrict the word-number channel to the
    # first body-scope numId encountered OUTSIDE the claims section so a
    # stray claim-list numId doesn't pollute the body numbering counter.
    body_num_id: str | None = None
    num_counters: dict[str, int] = {}
    in_claims = False
    claim_counter = 0

    for para in doc.paragraphs:
        text = _normalize_unicode(para.text).strip()
        if not text:
            continue

        # Resolve the paragraph's Word numId (None for unnumbered paragraphs).
        para_num_id = _get_paragraph_num_id(para)

        # Check for bracket headers to detect claims section boundary
        hm = _TW_BRACKET_RE.match(text)
        if hm:
            header = hm.group(1).strip()
            if header in _TW_CLAIMS_HEADERS:
                in_claims = True
                claim_counter = 0
                paragraphs.append(text)
                paragraph_word_numbers.append(None)
                continue
            elif in_claims:
                # Any other bracket header ends the claims section
                in_claims = False

        # Compute the body-scope 【NNNN】 label for this paragraph. Only a
        # single numId (the first encountered outside the claims section)
        # contributes to the public label channel; other numIds are ignored
        # for labeling purposes so claim-list numbering can't poison the
        # body counter.
        word_num: str | None = None
        if para_num_id is not None and not in_claims:
            if body_num_id is None:
                body_num_id = para_num_id
            if para_num_id == body_num_id:
                count = num_counters.get(para_num_id, 0) + 1
                num_counters[para_num_id] = count
                word_num = f"{count:04d}"

        if in_claims:
            # Check for Word numbering on this paragraph via the shared
            # helper (ADR-109 / Phase 8c): a truthy return means the
            # paragraph needs a synthetic "N. " prefix so claims_tw.py
            # can parse it.
            if _extract_numpr_claim_number(para) is not None:
                claim_counter += 1
                paragraphs.append(f"{claim_counter}. {text}")
            else:
                paragraphs.append(text)
            paragraph_word_numbers.append(None)
        else:
            paragraphs.append(text)
            paragraph_word_numbers.append(word_num)

    tracked_changes = detect_tracked_changes(doc)
    return LoadedTwDocument(
        paragraphs=paragraphs,
        paragraph_word_numbers=paragraph_word_numbers,
        has_tracked_changes=tracked_changes,
    )


def load_docx_cn(file_path: str | Path) -> LoadedCnDocument:
    """Load a CN patent .docx and return paragraphs grouped by Word section.

    The 五書模板 format uses Word section breaks to separate the five patent
    documents. Section titles appear in page headers (黑体 16pt centered),
    not in body text. This function reads those headers and groups body
    paragraphs by their containing Word section. Also detects tracked changes.
    """
    path = Path(file_path)
    if not path.exists():
        msg = f"File not found: {path}"
        raise FileNotFoundError(msg)
    if path.suffix.lower() != ".docx":
        msg = f"Not a .docx file: {path}"
        raise ValueError(msg)

    try:
        doc = Document(str(path))
    except Exception as exc:
        msg = f"Failed to open .docx: {exc}"
        raise ValueError(msg) from exc

    # Collect default header text for each Word section.
    # The 五書模板 uses different-first-page headers (first page empty,
    # default header has section title). We read the default header.
    section_headers: list[str] = []
    for section in doc.sections:
        header_text = "\n".join(
            _normalize_unicode(p.text).strip()
            for p in section.header.paragraphs
            if _normalize_unicode(p.text).strip()
        )
        section_headers.append(header_text)

    # Group paragraphs by Word section boundary.
    # Section breaks are marked by <w:sectPr> inside a paragraph's <w:pPr>.
    # The last section's <w:sectPr> is a direct child of <w:body> (no pPr).
    result: list[DocxSection] = []
    current_paras: list[str] = []
    current_numpr: list[bool] = []
    section_idx = 0

    for para in doc.paragraphs:
        text = _normalize_unicode(para.text).strip()
        if text:
            current_paras.append(text)
            # Parallel numPr flag — used by sections_cn to backfill
            # synthetic "N. " prefixes on claim paragraphs that lack a
            # typed number prefix (ADR-109).
            current_numpr.append(_extract_numpr_claim_number(para) is not None)

        # Check for section break in this paragraph
        pPr = para._element.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
            header = section_headers[section_idx] if section_idx < len(section_headers) else ""
            result.append(DocxSection(
                header_text=header,
                paragraphs=current_paras,
                numpr_flags=current_numpr,
            ))
            current_paras = []
            current_numpr = []
            section_idx += 1

    # Last section (body-level sectPr, not inside any paragraph)
    header = section_headers[section_idx] if section_idx < len(section_headers) else ""
    result.append(DocxSection(
        header_text=header,
        paragraphs=current_paras,
        numpr_flags=current_numpr,
    ))

    tracked_changes = detect_tracked_changes(doc)
    return LoadedCnDocument(sections=result, has_tracked_changes=tracked_changes)
